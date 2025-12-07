"""Resume formatting functions for DOCX, PDF, and text export"""
import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx_from_json(resume_data, filename="resume.docx"):
    """Generate a professional .docx file from structured resume JSON"""
    try:
        doc = Document()
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        header = resume_data.get('header', {})
        if header.get('name'):
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(header['name'])
            name_run.font.size = Pt(18)
            name_run.font.bold = True
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        contact_info = []
        if header.get('email'):
            contact_info.append(header['email'])
        if header.get('phone'):
            contact_info.append(header['phone'])
        if header.get('location'):
            contact_info.append(header['location'])
        if header.get('linkedin'):
            contact_info.append(header['linkedin'])
        if header.get('portfolio'):
            contact_info.append(header['portfolio'])
        
        if contact_info:
            contact_para = doc.add_paragraph(' | '.join(contact_info))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.runs[0].font.size = Pt(10)
        
        doc.add_paragraph()
        
        if header.get('title'):
            title_para = doc.add_paragraph(header['title'])
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.runs[0].font.size = Pt(12)
            title_para.runs[0].italic = True
            doc.add_paragraph()
        
        if resume_data.get('summary'):
            doc.add_heading('Professional Summary', level=2)
            summary_para = doc.add_paragraph(resume_data['summary'])
            summary_para.runs[0].font.size = Pt(11)
            doc.add_paragraph()
        
        skills = resume_data.get('skills_highlighted', [])
        if skills:
            doc.add_heading('Key Skills', level=2)
            skills_text = ' • '.join(skills)
            skills_para = doc.add_paragraph(skills_text)
            skills_para.runs[0].font.size = Pt(11)
            doc.add_paragraph()
        
        experience = resume_data.get('experience', [])
        if experience:
            doc.add_heading('Professional Experience', level=2)
            for exp in experience:
                exp_header = doc.add_paragraph()
                exp_header.add_run(exp.get('title', '')).bold = True
                if exp.get('company'):
                    exp_header.add_run(f" at {exp['company']}")
                if exp.get('dates'):
                    exp_header.add_run(f" | {exp['dates']}").italic = True
                
                bullets = exp.get('bullets', [])
                for bullet in bullets:
                    if bullet.strip():
                        bullet_para = doc.add_paragraph(bullet, style='List Bullet')
                        bullet_para.runs[0].font.size = Pt(10)
                
                doc.add_paragraph()
        
        if resume_data.get('education'):
            doc.add_heading('Education', level=2)
            edu_para = doc.add_paragraph(resume_data['education'])
            edu_para.runs[0].font.size = Pt(11)
            doc.add_paragraph()
        
        if resume_data.get('certifications'):
            doc.add_heading('Certifications & Awards', level=2)
            cert_para = doc.add_paragraph(resume_data['certifications'])
            cert_para.runs[0].font.size = Pt(11)
        
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io
        
    except Exception as e:
        st.error(f"Error generating DOCX: {e}")
        return None


def generate_pdf_from_json(resume_data, filename="resume.pdf"):
    """Generate a professional PDF file from structured resume JSON"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        
        pdf_io = BytesIO()
        doc = SimpleDocTemplate(pdf_io, pagesize=letter,
                               rightMargin=0.75*inch, leftMargin=0.75*inch,
                               topMargin=0.5*inch, bottomMargin=0.5*inch)
        
        elements = []
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='black',
            spaceAfter=6,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            textColor='black',
            spaceAfter=6,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            textColor='black',
            spaceAfter=6,
            leading=12
        )
        
        contact_style = ParagraphStyle(
            'CustomContact',
            parent=styles['Normal'],
            fontSize=9,
            textColor='black',
            spaceAfter=12,
            alignment=TA_CENTER
        )
        
        header = resume_data.get('header', {})
        if header.get('name'):
            elements.append(Paragraph(header['name'], title_style))
            elements.append(Spacer(1, 0.1*inch))
        
        contact_info = []
        if header.get('email'):
            contact_info.append(header['email'])
        if header.get('phone'):
            contact_info.append(header['phone'])
        if header.get('location'):
            contact_info.append(header['location'])
        if header.get('linkedin'):
            contact_info.append(header['linkedin'])
        if header.get('portfolio'):
            contact_info.append(header['portfolio'])
        
        if contact_info:
            elements.append(Paragraph(' | '.join(contact_info), contact_style))
            elements.append(Spacer(1, 0.1*inch))
        
        if header.get('title'):
            elements.append(Paragraph(header['title'], contact_style))
            elements.append(Spacer(1, 0.15*inch))
        
        if resume_data.get('summary'):
            elements.append(Paragraph('Professional Summary', heading_style))
            elements.append(Paragraph(resume_data['summary'], normal_style))
            elements.append(Spacer(1, 0.1*inch))
        
        skills = resume_data.get('skills_highlighted', [])
        if skills:
            elements.append(Paragraph('Key Skills', heading_style))
            skills_text = ' • '.join(skills)
            elements.append(Paragraph(skills_text, normal_style))
            elements.append(Spacer(1, 0.1*inch))
        
        experience = resume_data.get('experience', [])
        if experience:
            elements.append(Paragraph('Professional Experience', heading_style))
            for exp in experience:
                exp_header_parts = []
                if exp.get('title'):
                    exp_header_parts.append(f"<b>{exp['title']}</b>")
                if exp.get('company'):
                    exp_header_parts.append(f" at {exp['company']}")
                if exp.get('dates'):
                    exp_header_parts.append(f" | <i>{exp['dates']}</i>")
                
                if exp_header_parts:
                    elements.append(Paragraph(''.join(exp_header_parts), normal_style))
                
                bullets = exp.get('bullets', [])
                for bullet in bullets:
                    if bullet.strip():
                        elements.append(Paragraph(f"• {bullet}", normal_style))
                
                elements.append(Spacer(1, 0.1*inch))
        
        if resume_data.get('education'):
            elements.append(Paragraph('Education', heading_style))
            elements.append(Paragraph(resume_data['education'], normal_style))
            elements.append(Spacer(1, 0.1*inch))
        
        if resume_data.get('certifications'):
            elements.append(Paragraph('Certifications & Awards', heading_style))
            elements.append(Paragraph(resume_data['certifications'], normal_style))
        
        doc.build(elements)
        pdf_io.seek(0)
        return pdf_io
        
    except Exception as e:
        st.error(f"Error generating PDF: {e}")
        return None


def format_resume_as_text(resume_data):
    """Format structured resume JSON as plain text"""
    text = []
    
    header = resume_data.get('header', {})
    if header.get('name'):
        text.append(header['name'].upper())
        text.append("")
    
    contact = []
    if header.get('email'):
        contact.append(header['email'])
    if header.get('phone'):
        contact.append(header['phone'])
    if header.get('location'):
        contact.append(header['location'])
    if header.get('linkedin'):
        contact.append(header['linkedin'])
    if header.get('portfolio'):
        contact.append(header['portfolio'])
    
    if contact:
        text.append(' | '.join(contact))
        text.append("")
    
    if header.get('title'):
        text.append(header['title'])
        text.append("")
    
    if resume_data.get('summary'):
        text.append("PROFESSIONAL SUMMARY")
        text.append("-" * 50)
        text.append(resume_data['summary'])
        text.append("")
    
    skills = resume_data.get('skills_highlighted', [])
    if skills:
        text.append("KEY SKILLS")
        text.append("-" * 50)
        text.append(' • '.join(skills))
        text.append("")
    
    experience = resume_data.get('experience', [])
    if experience:
        text.append("PROFESSIONAL EXPERIENCE")
        text.append("-" * 50)
        for exp in experience:
            exp_line = exp.get('title', '')
            if exp.get('company'):
                exp_line += f" at {exp['company']}"
            if exp.get('dates'):
                exp_line += f" | {exp['dates']}"
            text.append(exp_line)
            
            bullets = exp.get('bullets', [])
            for bullet in bullets:
                if bullet.strip():
                    text.append(f"  • {bullet}")
            text.append("")
    
    if resume_data.get('education'):
        text.append("EDUCATION")
        text.append("-" * 50)
        text.append(resume_data['education'])
        text.append("")
    
    if resume_data.get('certifications'):
        text.append("CERTIFICATIONS & AWARDS")
        text.append("-" * 50)
        text.append(resume_data['certifications'])
    
    return '\n'.join(text)
