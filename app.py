import warnings
import os
import gc
import sys
warnings.filterwarnings('ignore')

# Streamlit Cloud optimization: Set log level to error for production
os.environ['STREAMLIT_LOG_LEVEL'] = 'error'
# Streamlit Cloud: Disable telemetry
os.environ['STREAMLIT_BROWSER_GATHER_USAGE_STATS'] = 'false'
# Reduce SQLite memory footprint (used by ChromaDB)
os.environ['SQLITE_TMPDIR'] = '/tmp'
# Disable ChromaDB anonymous telemetry
os.environ['ANONYMIZED_TELEMETRY'] = 'false'
# Streamlit Cloud memory optimization: Limit SQLite cache size
os.environ['SQLITE_DEFAULT_PAGE_SIZE'] = '1024'

import streamlit as st

# Set page config FIRST (must be first Streamlit command)
st.set_page_config(
    page_title="CareerLens - AI Career Intelligence",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'About': "CareerLens - AI-powered career intelligence platform"
    }
)

# Core imports with error handling
import requests
import numpy as np
import pandas as pd
from sklearn.metrics.pairwise import cosine_similarity
import time
import math
from datetime import datetime, timedelta, timezone
from email.utils import parsedate_to_datetime
import json
import re
from io import BytesIO
import hashlib
import base64

# Document processing imports with graceful fallback
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    st.warning("⚠️ PyPDF2 not available - PDF upload disabled")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("⚠️ python-docx not available - DOCX upload disabled")

# ChromaDB import with graceful fallback (memory-intensive)
CHROMADB_AVAILABLE = False
chromadb = None
Settings = None
try:
    import chromadb as _chromadb
    from chromadb.config import Settings as _Settings
    chromadb = _chromadb
    Settings = _Settings
    CHROMADB_AVAILABLE = True
except ImportError as e:
    # ChromaDB not available - will use simple in-memory storage
    pass
except Exception as e:
    # ChromaDB import failed - will use simple in-memory storage
    pass

# Tokenization import
try:
    import tiktoken
    TIKTOKEN_AVAILABLE = True
except ImportError:
    TIKTOKEN_AVAILABLE = False

# Periodic garbage collection to reduce memory pressure on Streamlit Cloud
gc.collect()


# -----------------------------------------------------------------------------
# MEMORY MANAGEMENT UTILITIES
# -----------------------------------------------------------------------------
def _cleanup_session_state():
    """Clean up old/stale data from session state to prevent memory bloat.
    
    Called periodically to ensure session state doesn't grow unbounded.
    """
    # Limit cache sizes to prevent memory issues
    MAX_CACHE_ENTRIES = 10
    MAX_SKILL_CACHE_SIZE = 500
    
    # Clean up jobs cache - keep only most recent entries
    if 'jobs_cache' in st.session_state and isinstance(st.session_state.jobs_cache, dict):
        cache = st.session_state.jobs_cache
        if len(cache) > MAX_CACHE_ENTRIES:
            # Sort by timestamp and keep only the newest entries
            sorted_keys = sorted(
                cache.keys(),
                key=lambda k: cache[k].get('timestamp', ''),
                reverse=True
            )
            keys_to_remove = sorted_keys[MAX_CACHE_ENTRIES:]
            for key in keys_to_remove:
                del cache[key]
    
    # Clean up skill embeddings cache
    if 'skill_embeddings_cache' in st.session_state:
        cache = st.session_state.skill_embeddings_cache
        if len(cache) > MAX_SKILL_CACHE_SIZE:
            # Keep only half of the entries (simple LRU-like cleanup)
            keys_to_remove = list(cache.keys())[:-MAX_SKILL_CACHE_SIZE//2]
            for key in keys_to_remove:
                del cache[key]
    
    if 'user_skills_embeddings_cache' in st.session_state:
        cache = st.session_state.user_skills_embeddings_cache
        if len(cache) > MAX_SKILL_CACHE_SIZE:
            keys_to_remove = list(cache.keys())[:-MAX_SKILL_CACHE_SIZE//2]
            for key in keys_to_remove:
                del cache[key]
    
    # Force garbage collection after cleanup
    gc.collect()


# Run cleanup on module load (each session)
try:
    _cleanup_session_state()
except Exception:
    pass  # Ignore errors during cleanup (session state might not be initialized yet)


# -----------------------------------------------------------------------------
# HELPER FUNCTION: LOAD IMAGE AS BASE64 FOR HERO BANNER
# -----------------------------------------------------------------------------
def get_img_as_base64(file):
    """Convert an image file to base64 string for embedding in HTML"""
    with open(file, "rb") as f:
        data = f.read()
    return base64.b64encode(data).decode()


def _coerce_positive_int(value, default, minimum=1):
    """Ensure configuration values are positive integers."""
    try:
        if value is None:
            raise ValueError
        coerced = int(value)
        return max(minimum, coerced)
    except (ValueError, TypeError):
        return default


def _coerce_positive_float(value, default, minimum=0.0):
    """Ensure configuration values are positive floats."""
    try:
        if value is None:
            raise ValueError
        coerced = float(value)
        return max(minimum, coerced)
    except (ValueError, TypeError):
        return default


def _get_config_int(key, default, minimum=1):
    """Look up configuration values from Streamlit secrets or environment."""
    try:
        # Safely access secrets - handle case where Streamlit might not be fully initialized
        secrets_value = st.secrets.get(key)
    except (AttributeError, RuntimeError, KeyError, Exception):
        # Handle case where Streamlit isn't initialized yet, secrets aren't available, or key doesn't exist
        secrets_value = None
    env_value = os.getenv(key)
    candidate = secrets_value if secrets_value not in (None, "") else env_value
    return _coerce_positive_int(candidate, default, minimum)


def _get_config_float(key, default, minimum=0.0):
    """Look up float configuration values from Streamlit secrets or environment."""
    try:
        secrets_value = st.secrets.get(key)
    except (AttributeError, RuntimeError, KeyError, Exception):
        secrets_value = None
    env_value = os.getenv(key)
    candidate = secrets_value if secrets_value not in (None, "") else env_value
    return _coerce_positive_float(candidate, default, minimum)


# Reduced default batch size to avoid rate limits (can be overridden via env var)
# Note: These are evaluated at module import time, but _get_config_int handles errors gracefully
DEFAULT_EMBEDDING_BATCH_SIZE = _get_config_int("EMBEDDING_BATCH_SIZE", 15, minimum=5)  # Reduced from 20
DEFAULT_MAX_JOBS_TO_INDEX = _get_config_int("MAX_JOBS_TO_INDEX", 25, minimum=10)  # Reduced from 50 to minimize API calls
# Rate limiting: delay between batches in seconds (gentle spacing between successful batches)
# Note: api_call_with_retry() handles 429 errors with exponential backoff automatically
EMBEDDING_BATCH_DELAY = _get_config_float("EMBEDDING_BATCH_DELAY", 0.5, minimum=0.0)  # Reduced from 1s to 0.5s
# RapidAPI rate limiting: max requests per minute (default: 3 for free tier)
RAPIDAPI_MAX_REQUESTS_PER_MINUTE = _get_config_int("RAPIDAPI_MAX_REQUESTS_PER_MINUTE", 3, minimum=1)
# Profile extraction: skip Pass 2 (self-correction) by default for faster processing
# Set to True to enable the self-correction pass (adds 10-30s but improves accuracy)
ENABLE_PROFILE_PASS2 = os.getenv("ENABLE_PROFILE_PASS2", "false").lower() in ("true", "1", "yes")


def _determine_index_limit(total_jobs, desired_top_matches):
    """Cap how many jobs we embed per search to avoid unnecessary API calls.
    
    Optimized to reduce API calls while maintaining good match quality.
    For 15 desired matches, we index ~20-25 jobs (not 45-50 as before).
    """
    # More conservative: 1.5x desired matches instead of 3x
    baseline = max(desired_top_matches + 10, 15)
    limit = min(DEFAULT_MAX_JOBS_TO_INDEX, baseline)
    return min(total_jobs, limit)


def _parse_retry_after_value(value):
    """Convert Retry-After style header values into seconds."""
    if not value:
        return None
    value = value.strip()
    if not value:
        return None
    # Try numeric (int or float)
    try:
        seconds = float(value)
        if seconds >= 0:
            return int(math.ceil(seconds))
    except (ValueError, TypeError):
        pass
    # Try HH:MM:SS format
    if value.count(':') == 2:
        try:
            hours, minutes, seconds = value.split(':')
            total_seconds = int(hours) * 3600 + int(minutes) * 60 + int(float(seconds))
            if total_seconds >= 0:
                return total_seconds
        except (ValueError, TypeError):
            pass
    # Try HTTP-date format
    try:
        retry_time = parsedate_to_datetime(value)
        if retry_time:
            if retry_time.tzinfo is None:
                retry_time = retry_time.replace(tzinfo=timezone.utc)
            now = datetime.now(timezone.utc)
            delta = (retry_time - now).total_seconds()
            if delta > 0:
                return int(math.ceil(delta))
    except (TypeError, ValueError, OverflowError):
        pass
    return None


def _extract_delay_from_body(response):
    """Attempt to read retry hints from JSON/text error bodies."""
    if response is None:
        return None
    message = None
    try:
        data = response.json()
        if isinstance(data, dict):
            error = data.get('error') or {}
            if isinstance(error, dict):
                message = error.get('message') or error.get('code')
            if not message:
                message = data.get('message')
    except (ValueError, json.JSONDecodeError):
        pass
    if not message:
        message = response.text or ""
    if not message:
        return None
    match = re.search(r'after\s+(\d+)\s+seconds?', message, re.IGNORECASE)
    if match:
        try:
            seconds = int(match.group(1))
            if seconds >= 0:
                return seconds
        except ValueError:
            pass
    return None


def _determine_retry_delay(response, fallback_delay, max_delay):
    """Use headers/body hints to determine how long to wait before retrying."""
    if response is not None:
        headers = response.headers or {}
        header_candidates = [
            'Retry-After',
            'x-ms-retry-after-ms',
            'x-ms-retry-after',
            'x-ratelimit-reset-requests',
            'x-ratelimit-reset-tokens',
            'x-ratelimit-reset',
        ]
        for header in header_candidates:
            raw_value = headers.get(header)
            if not raw_value:
                continue
            # Headers with milliseconds suffix
            if header.endswith('-ms'):
                try:
                    ms = float(raw_value)
                    if ms >= 0:
                        seconds = int(math.ceil(ms / 1000.0))
                        return max(1, min(seconds, max_delay)), f"header:{header}"
                except (ValueError, TypeError):
                    continue
            else:
                parsed = _parse_retry_after_value(raw_value)
                if parsed is not None:
                    return max(1, min(parsed, max_delay)), f"header:{header}"
        # Fallback to error body hints
        body_delay = _extract_delay_from_body(response)
        if body_delay is not None:
            return max(1, min(body_delay, max_delay)), "body"
    return max(1, min(fallback_delay, max_delay)), "fallback"


def _calculate_exponential_delay(initial_delay, attempt, max_delay):
    """Calculate exponential backoff delay for the current retry attempt."""
    return max(1, min(initial_delay * (2 ** attempt), max_delay))


def _chunked_sleep(delay, message_prefix=""):
    """Sleep in small chunks to prevent WebSocket timeout on Streamlit Cloud.
    
    Streamlit Cloud WebSocket connections can timeout if there's no activity for
    extended periods. This function breaks long sleeps into smaller chunks with
    UI updates to keep the connection alive.
    
    IMPORTANT: Streamlit Cloud's logstream WebSocket can disconnect after ~15-20s
    of inactivity. We use 2-second chunks to ensure frequent UI updates.
    """
    if delay <= 2:
        time.sleep(delay)
        return
    
    status_placeholder = st.empty()
    remaining = int(delay)
    while remaining > 0:
        if message_prefix:
            status_placeholder.caption(f"{message_prefix} ({remaining}s remaining...)")
        else:
            # Even without a message, update UI to keep WebSocket alive
            status_placeholder.caption(f"⏳ Processing... ({remaining}s)")
        chunk = min(2, remaining)  # Max 2 seconds at a time for WebSocket stability
        time.sleep(chunk)
        remaining -= chunk
    status_placeholder.empty()


def _websocket_keepalive(message=None):
    """Send a lightweight UI update to keep WebSocket connection alive.
    
    Call this periodically during long-running operations that don't have
    their own progress indicators.
    
    The Streamlit Cloud logstream WebSocket (/~/logstream) can disconnect
    after ~15-20s of inactivity. This function sends a minimal UI update
    to keep the connection alive.
    
    Args:
        message: Optional status message to display briefly (then cleared)
    """
    # Use an empty placeholder to trigger a minimal UI update
    placeholder = st.empty()
    if message:
        placeholder.caption(f"⏳ {message}")
        # Brief display before clearing
        time.sleep(0.1)
    placeholder.empty()


def api_call_with_retry(func, max_retries=3, initial_delay=1, max_delay=60):
    """
    Execute an API call with exponential backoff retry logic for rate limit errors (429).
    
    Args:
        func: Function that makes the API call and returns a requests.Response object
        max_retries: Maximum number of retry attempts
        initial_delay: Initial delay in seconds before first retry
        max_delay: Maximum delay in seconds between retries
    
    Returns:
        Response object if successful, None otherwise
    """
    for attempt in range(max_retries):
        try:
            response = func()
            
            # Success case
            if response.status_code in [200, 201]:
                return response
            
            # Rate limit error (429) - retry with exponential backoff
            elif response.status_code == 429:
                if attempt < max_retries - 1:
                    fallback_delay = _calculate_exponential_delay(initial_delay, attempt, max_delay)
                    delay, delay_source = _determine_retry_delay(response, fallback_delay, max_delay)
                    source_note = ""
                    if delay_source != "fallback":
                        source_note = f" (server hint: {delay_source})"
                    # Only show detailed message on first retry, use less intrusive indicator for subsequent retries
                    if attempt == 0:
                        st.warning(
                            f"⏳ Rate limit reached. Retrying in {delay} seconds{source_note}... "
                            f"(Attempt {attempt + 1}/{max_retries})"
                        )
                    else:
                        # Use caption for subsequent retries (less intrusive)
                        st.caption(f"⏳ Retrying... ({attempt + 1}/{max_retries})")
                    # Use chunked sleep to prevent WebSocket timeout
                    _chunked_sleep(delay, f"⏳ Retry {attempt + 1}/{max_retries}")
                    continue
                else:
                    # Max retries exceeded
                    error_msg = (
                        "🚫 **Rate Limit Exceeded**\n\n"
                        "The API rate limit has been reached. Please:\n"
                        "1. Wait a few minutes and try again\n"
                        "2. Reduce the number of jobs you're searching for\n"
                        "3. Check your API quota/limits\n\n"
                        f"Status: {response.status_code}"
                    )
                    st.error(error_msg)
                    return None
            
            # Other HTTP errors - don't retry
            else:
                return response
                
        except requests.exceptions.Timeout:
            if attempt < max_retries - 1:
                delay = _calculate_exponential_delay(initial_delay, attempt, max_delay)
                st.warning(f"⏳ Request timed out. Retrying in {delay} seconds... (Attempt {attempt + 1}/{max_retries})")
                _chunked_sleep(delay)
                continue
            else:
                st.error("❌ Request timed out after multiple attempts. Please try again later.")
                return None
        
        except requests.exceptions.RequestException as e:
            if attempt < max_retries - 1:
                delay = _calculate_exponential_delay(initial_delay, attempt, max_delay)
                st.warning(f"⏳ Network error. Retrying in {delay} seconds... (Attempt {attempt + 1}/{max_retries})")
                _chunked_sleep(delay)
                continue
            else:
                st.error(f"❌ Network error after multiple attempts: {e}")
                return None
        
        except Exception as e:
            # Unexpected errors - don't retry
            st.error(f"❌ Unexpected error: {e}")
            return None
    
    return None

# Note: st.set_page_config is called at the top of the file (after imports)

# -----------------------------------------------------------------------------
# LOGO LOADING FOR HERO BANNER (Base64 for CSS/HTML injection)
# -----------------------------------------------------------------------------
# Try to load the logo. If it doesn't exist, we'll just use a placeholder.
# This ensures the app doesn't crash if the file is missing.
# We load as Base64 so we can inject it into HTML/CSS for the watermark effect.
_logo_base64 = ""
_logo_html = ""

# Check for logo.png first, then fallback to CareerLens_Logo.png
logo_paths = ["logo.png", "CareerLens_Logo.png"]
for logo_path in logo_paths:
    if os.path.exists(logo_path):
        try:
            _logo_base64 = get_img_as_base64(logo_path)
            _logo_html = f'<img src="data:image/png;base64,{_logo_base64}" class="hero-bg-logo">'
            break
        except Exception as e:
            _logo_html = '<div class="hero-bg-logo"></div>'
            break
else:
    # No logo found - use empty placeholder
    _logo_html = '<div class="hero-bg-logo"></div>'

# Theme will be automatically detected from system preferences via JavaScript
# CSS uses attribute selectors to respond to data-theme attribute

st.markdown("""
<style>
    /* CareerLens Design System - CSS Variables */
    /* New Design Color Scheme */
    :root {{
        /* New design primary colors */
        --navy: #0f172a;
        --cyan: #00d2ff;
        --bg-gray: #f3f4f6;
        
        /* Existing design system colors */
        --primary-accent: #0F62FE;
        --action-accent: #0F62FE;
        --bg-main: #f3f4f6;
        --bg-container: #F4F7FC;
        --card-bg: #FFFFFF;
        --text-primary: #161616;
        --text-secondary: #161616;
        --text-muted: #161616;
        --border-color: #E0E0E0;
        --hover-bg: #F0F0F0;
        --success-green: #10B981;
        --warning-amber: #F59E0B;
        --error-red: #EF4444;
        --navy-deep: #1e3a5f;
        --navy-light: #2C3E50;
    }}
    
    /* Dark mode override */
    [data-theme="dark"],
    html[data-theme="dark"],
    html[data-theme="dark"] :root {{
        --primary-accent: #4589FF;
        --action-accent: #4589FF;
        --bg-main: #161616;
        --bg-container: #262626;
        --card-bg: #262626;
        --text-primary: #F4F4F4;
        --text-secondary: #F4F4F4;
        --text-muted: #F4F4F4;
        --border-color: #3D3D3D;
        --hover-bg: #333333;
        --navy: #1e293b;
        --cyan: #22d3ee;
        --bg-gray: #1f2937;
    }}
    
    /* --- HIDE DEFAULT STREAMLIT ELEMENTS (Header, Footer, Menu) --- */
    #MainMenu {{visibility: hidden;}}
    footer {{visibility: hidden;}}
    header[data-testid="stHeader"] {{visibility: hidden; height: 0; padding: 0; margin: 0;}}
    .stDeployButton {{display: none;}}
    
    /* Global Styling */
    .stApp {{
        background-color: var(--bg-gray);
        color: var(--text-primary);
    }}
    
    /* --- SIDEBAR STYLING (Navy Theme) --- */
    [data-testid="stSidebar"] {{
        background-color: var(--navy);
        padding: 2rem 1rem;
    }}
    [data-testid="stSidebar"] * {{
        color: #94a3b8 !important;
    }}
    [data-testid="stSidebar"] h1,
    [data-testid="stSidebar"] h2,
    [data-testid="stSidebar"] h3,
    [data-testid="stSidebar"] .stMarkdown h2,
    [data-testid="stSidebar"] .stMarkdown h3 {{
        color: white !important;
    }}
    [data-testid="stSidebar"] .stRadio label:hover,
    [data-testid="stSidebar"] .stSelectbox label:hover {{
        color: white !important;
    }}
    [data-testid="stSidebar"] .stButton > button {{
        background-color: var(--cyan) !important;
        color: var(--navy) !important;
        font-weight: 600 !important;
    }}
    [data-testid="stSidebar"] .stButton > button:hover {{
        background-color: #06b6d4 !important;
    }}
    
    /* --- HERO BANNER --- */
    .hero-container {{
        background: linear-gradient(135deg, var(--navy) 0%, #112545 100%);
        padding: 40px;
        border-radius: 12px;
        color: white;
        position: relative;
        overflow: hidden;
        margin-bottom: 30px;
        box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1);
    }}
    .hero-content {{
        position: relative;
        z-index: 2;
    }}
    .hero-title {{
        font-size: 32px;
        font-weight: 700;
        margin: 0;
        color: white;
    }}
    .hero-subtitle {{
        color: #94a3b8;
        font-size: 16px;
        margin-top: 10px;
    }}
    .hero-bg-logo {{
        position: absolute;
        right: -30px;
        top: -30px;
        width: 250px;
        opacity: 0.15;
        transform: rotate(-15deg);
        pointer-events: none;
        z-index: 1;
    }}
    
    /* --- DASHBOARD METRIC CARDS --- */
    .dashboard-metric-card {{
        background: white;
        padding: 20px;
        border-radius: 8px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        text-align: center;
    }}
    .dashboard-metric-label {{
        font-size: 12px;
        color: #6b7280;
        text-transform: uppercase;
        letter-spacing: 1px;
    }}
    .dashboard-metric-value {{
        font-size: 28px;
        font-weight: 700;
        color: #111827;
        margin-top: 5px;
    }}
    
    /* Dark mode hero adjustments */
    [data-theme="dark"] .hero-container {{
        background: linear-gradient(135deg, #1e293b 0%, #0f172a 100%);
    }}
    [data-theme="dark"] .dashboard-metric-card {{
        background: var(--card-bg);
    }}
    [data-theme="dark"] .dashboard-metric-value {{
        color: var(--text-primary);
    }}
    
    /* Card Styling */
    .metric-card {{
        background-color: var(--card-bg);
        border: 1px solid var(--border-color);
        border-radius: 12px;
        padding: 24px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        margin-bottom: 16px;
    }}
    
    /* Button Styling */
    .stButton > button {{
        background-color: var(--primary-accent);
        color: white;
        border-radius: 8px;
        border: none;
        padding: 0.5rem 2rem;
        font-weight: 500;
        transition: all 0.3s ease;
    }}
    
    .stButton > button:hover {{
        background-color: var(--primary-accent);
        opacity: 0.8;
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(15, 98, 254, 0.3);
    }}
    
    /* Secondary Button */
    .secondary-button {{
        background-color: transparent;
        color: var(--primary-accent);
        border: 1px solid var(--primary-accent);
        border-radius: 8px;
        padding: 0.5rem 1rem;
        font-weight: 500;
        transition: all 0.3s ease;
    }}
    
    .secondary-button:hover {{
        background-color: var(--primary-accent);
        color: white;
    }}
    
    /* Table Styling */
    .dataframe {{
        background-color: var(--card-bg);
        border: 1px solid var(--border-color);
        border-radius: 12px;
        overflow: hidden;
    }}
    
    .dataframe th {{
        background-color: var(--bg-container);
        color: var(--text-primary);
        font-weight: 600;
        padding: 12px;
        border-bottom: 2px solid var(--border-color);
    }}
    
    .dataframe td {{
        padding: 12px;
        border-bottom: 1px solid var(--border-color);
    }}
    
    .dataframe tr:hover {{
        background-color: var(--hover-bg);
    }}
    
    /* Progress Bar */
    .progress-bar {{
        background-color: var(--border-color);
        border-radius: 10px;
        height: 8px;
        overflow: hidden;
    }}
    
    .progress-fill {{
        background-color: var(--primary-accent);
        height: 100%;
        border-radius: 10px;
        transition: width 0.3s ease;
    }}
    
    /* Upload Section */
    [data-testid="stFileUploader"] {{
        background-color: var(--card-bg);
        border: 2px dashed var(--border-color);
        border-radius: 12px;
        padding: 2rem;
    }}
    
    /* Slider Styling */
    .stSlider {{
        padding: 1rem 0;
    }}
    
    /* Header Styling */
    h1, h2, h3 {{
        color: var(--text-primary);
    }}
    
    h1 {{
        font-size: 36px;
        font-weight: 700;
        margin-bottom: 2rem;
    }}
    
    h2 {{
        font-size: 28px;
        font-weight: 600;
        margin-bottom: 1.5rem;
    }}
    
    /* Metric Container */
    .metric-container {{
        display: flex;
        gap: 20px;
        margin-bottom: 2rem;
    }}
    
    .metric-label {{
        font-size: 14px;
        color: #6B7280;
        margin-bottom: 8px;
        font-weight: 500;
    }}
    
    .metric-value {{
        font-size: 28px;
        font-weight: 700;
        color: var(--text-primary);
    }}
    
    /* Expandable Section */
    .expander-content {{
        background-color: var(--bg-container);
        border-left: 4px solid var(--primary-accent);
        padding: 16px;
        margin: 8px 0;
        border-radius: 4px;
    }}
    
    /* Input Fields */
    .stTextInput > div > div > input,
    .stSelectbox > div > div > select {{
        background-color: var(--card-bg);
        color: var(--text-primary);
        border: 1px solid var(--border-color);
        border-radius: 8px;
    }}
    
    /* Multi-select */
    .stMultiSelect {{
        background-color: var(--card-bg);
        border-radius: 8px;
    }}
    
    /* Badge/Tag Styling */
    .fit-type-badge {{
        background-color: var(--bg-container);
        color: var(--text-primary);
        padding: 4px 12px;
        border-radius: 16px;
        font-size: 12px;
        font-weight: 500;
        display: inline-block;
    }}
    
    /* Logo Section */
    .logo-container {{
        display: flex;
        align-items: center;
        gap: 12px;
        margin-bottom: 2rem;
    }}
    
    .logo-text {{
        font-size: 24px;
        font-weight: 700;
        color: var(--text-primary);
    }}
    
    /* Divider */
    hr {{
        border: none;
        border-top: 1px solid var(--border-color);
        margin: 2rem 0;
    }}
    
    /* Main Header */
    .main-header {{
        font-size: 3rem;
        font-weight: bold;
        color: var(--primary-accent);
        text-align: center;
        margin-bottom: 1rem;
        letter-spacing: -0.02em;
    }}
    
    .sub-header {{
        font-size: 1.2rem;
        text-align: center;
        color: var(--text-secondary);
        margin-bottom: 2rem;
    }}
    
    /* Step Hierarchy - Visual Markers */
    .step-marker {{
        display: inline-flex;
        align-items: center;
        justify-content: center;
        width: 2.5rem;
        height: 2.5rem;
        border-radius: 50%;
        background-color: var(--primary-accent);
        color: white;
        font-size: 1.2rem;
        font-weight: bold;
        margin-right: 1rem;
        flex-shrink: 0;
    }}
    
    .step-container {{
        margin: 2rem 0;
        padding: 2rem;
        background-color: var(--bg-container);
        border-radius: 12px;
        border: none;
    }}
    
    /* Job Card - Minimalist Design */
    .job-card {{
        background-color: var(--bg-container);
        padding: 1.5rem;
        border-radius: 12px;
        margin-bottom: 1.5rem;
        border: none;
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }}
    
    .job-card:hover {{
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
    }}
    
    /* Match Score - Action Accent */
    .match-score {{
        background-color: var(--action-accent);
        color: white;
        padding: 0.4rem 1rem;
        border-radius: 20px;
        font-weight: bold;
        display: inline-block;
        font-size: 0.9rem;
    }}
    
    /* Tags */
    .tag {{
        display: inline-block;
        background-color: var(--bg-container);
        color: var(--text-primary);
        padding: 0.3rem 0.8rem;
        border-radius: 12px;
        margin: 0.2rem;
        font-size: 0.85rem;
        border: none;
    }}
    
    /* Borderless Input Fields */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea,
    .stSelectbox > div > div > select {{
        border: none !important;
        border-bottom: 1px solid rgba(0, 0, 0, 0.1) !important;
        border-radius: 0 !important;
        background-color: transparent !important;
        padding: 0.5rem 0 !important;
        transition: border-color 0.3s ease !important;
    }}
    
    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus {{
        border-bottom: 2px solid var(--primary-accent) !important;
        outline: none !important;
        box-shadow: none !important;
    }}
    
    /* Dark mode input styling */
    [data-theme="dark"] .stTextInput > div > div > input,
    [data-theme="dark"] .stTextArea > div > div > textarea,
    [data-theme="dark"] .stSelectbox > div > div > select {{
        border-bottom: 1px solid rgba(255, 255, 255, 0.2) !important;
        color: var(--text-primary) !important;
    }}
    
    [data-theme="dark"] .stTextInput > div > div > input:focus,
    [data-theme="dark"] .stTextArea > div > div > textarea:focus {{
        border-bottom: 2px solid var(--primary-accent) !important;
    }}
    
    /* Buttons - Rounded with Primary Accent */
    .stButton > button {{
        border-radius: 8px !important;
        border: none !important;
        font-weight: 500 !important;
        transition: all 0.2s ease !important;
    }}
    
    .stButton > button[kind="primary"] {{
        background-color: var(--primary-accent) !important;
        color: white !important;
    }}
    
    .stButton > button[kind="primary"]:hover {{
        background-color: var(--primary-accent) !important;
        opacity: 0.9 !important;
        transform: translateY(-1px) !important;
    }}
    
    /* Sliders - Elegant Design */
    /* Slider track background - subtle gray */
    .stSlider [data-baseweb="slider"] > div:first-child {{
        background-color: rgba(0, 0, 0, 0.08) !important;
    }}
    
    [data-theme="dark"] .stSlider [data-baseweb="slider"] > div:first-child {{
        background-color: rgba(255, 255, 255, 0.15) !important;
    }}
    
    /* Slider track fill - primary accent color */
    .stSlider [data-baseweb="slider"] > div > div {{
        background-color: var(--primary-accent) !important;
    }}
    
    /* Target slider track background */
    .stSlider [data-baseweb="slider-track"] {{
        background-color: rgba(0, 0, 0, 0.08) !important;
    }}
    
    [data-theme="dark"] .stSlider [data-baseweb="slider-track"] {{
        background-color: rgba(255, 255, 255, 0.15) !important;
    }}
    
    /* Slider thumb/handle - primary accent with hover effect */
    .stSlider [data-baseweb="slider-thumb"],
    .stSlider [role="slider"] {{
        background-color: var(--primary-accent) !important;
        border-color: var(--primary-accent) !important;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1) !important;
        transition: transform 0.2s ease, box-shadow 0.2s ease !important;
    }}
    
    .stSlider [data-baseweb="slider-thumb"]:hover,
    .stSlider [role="slider"]:hover {{
        transform: scale(1.1) !important;
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.15) !important;
    }}
    
    /* Slider value display */
    .stSlider label {{
        color: var(--text-primary) !important;
    }}
    
    /* Status Indicators - Profile Ready */
    .profile-ready {{
        display: inline-flex;
        align-items: center;
        gap: 0.5rem;
        color: var(--success-green);
        font-weight: 500;
    }}
    
    .profile-ready::before {{
        content: "✓";
        display: inline-flex;
        align-items: center;
        justify-content: center;
        width: 1.2rem;
        height: 1.2rem;
        border-radius: 50%;
        background-color: var(--success-green);
        color: white;
        font-size: 0.8rem;
        font-weight: bold;
    }}
    
    /* Match Score Display - Prominent with Action Accent */
    .match-score-display {{
        font-size: 2rem;
        font-weight: bold;
        color: var(--action-accent);
        text-align: center;
    }}
    
    /* Step 3 Containers - Automatic/Manual Search */
    .step3-container {{
        background-color: var(--bg-container);
        padding: 1.5rem;
        border-radius: 12px;
        margin: 1rem 0;
    }}
    
    /* Alternative: Style columns directly for Step 3 */
    div[data-testid="column"]:has(.step3-container) {{
        padding: 0.5rem;
    }}
    
    /* Ensure container background shows through */
    .step3-container > * {{
        position: relative;
        z-index: 1;
    }}
    
    /* Minimal Dividers */
    hr {{
        border: none;
        border-top: 1px solid rgba(0, 0, 0, 0.08);
        margin: 2rem 0;
    }}
    
    [data-theme="dark"] hr {{
        border-top: 1px solid rgba(255, 255, 255, 0.1);
    }}
    
    /* Remove default Streamlit borders and shadows */
    .element-container {{
        border: none !important;
        box-shadow: none !important;
    }}
    
    /* Card-like containers for Step 3 */
    .matching-container {{
        background-color: var(--bg-container);
        padding: 1.5rem;
        border-radius: 12px;
        margin-bottom: 1rem;
    }}
    
    /* Ensure proper spacing and layout */
    .main .block-container {{
        padding-top: 2rem;
        padding-bottom: 2rem;
    }}
    
    /* Improve form styling */
    .stForm {{
        border: none;
        background-color: transparent;
    }}
    
    /* File uploader styling */
    .stFileUploader > div {{
        border: 1px dashed rgba(0, 0, 0, 0.2);
        border-radius: 8px;
        padding: 1rem;
        background-color: var(--bg-container);
    }}
    
    [data-theme="dark"] .stFileUploader > div {{
        border-color: rgba(255, 255, 255, 0.2);
    }}
    
    /* Expander styling */
    .streamlit-expanderHeader {{
        color: var(--text-primary);
        font-weight: 500;
    }}
    
    /* Metric styling */
    [data-testid="stMetricValue"] {{
        color: var(--text-primary);
    }}
    
    [data-testid="stMetricLabel"] {{
        color: var(--text-secondary);
    }}
    
    /* Info/Warning/Success boxes */
    .stAlert {{
        border-radius: 8px;
        border: none;
    }}
    
    /* Number input styling */
    .stNumberInput > div > div > input {{
        border: none !important;
        border-bottom: 1px solid rgba(0, 0, 0, 0.1) !important;
        border-radius: 0 !important;
        background-color: transparent !important;
    }}
    
    [data-theme="dark"] .stNumberInput > div > div > input {{
        border-bottom: 1px solid rgba(255, 255, 255, 0.2) !important;
        color: var(--text-primary) !important;
    }}
    
    .stNumberInput > div > div > input:focus {{
        border-bottom: 2px solid var(--primary-accent) !important;
        outline: none !important;
    }}
    
    /* Executive Dashboard Styles */
    .dashboard-header {{
        font-size: 2.5rem;
        font-weight: 700;
        color: var(--primary-accent);
        margin-bottom: 0.5rem;
        letter-spacing: -0.02em;
    }}
    
    .dashboard-subtitle {{
        font-size: 1rem;
        color: var(--text-secondary);
        margin-bottom: 2rem;
    }}
    
    .metric-card {{
        background-color: var(--bg-container);
        padding: 1.5rem;
        border-radius: 12px;
        border: 1px solid rgba(0, 0, 0, 0.05);
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }}
    
    .metric-card:hover {{
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
    }}
    
    .metric-value {{
        font-size: 2rem;
        font-weight: bold;
        color: var(--primary-accent);
        margin: 0.5rem 0;
    }}
    
    .metric-label {{
        font-size: 0.9rem;
        color: var(--text-secondary);
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }}
    
    .job-table-row {{
        cursor: pointer;
        transition: background-color 0.2s ease;
        border-bottom: 1px solid rgba(0, 0, 0, 0.05);
    }}
    
    .job-table-row:hover {{
        background-color: var(--bg-container);
    }}
    
    .match-score-badge {{
        display: inline-flex;
        align-items: center;
        padding: 0.4rem 0.8rem;
        border-radius: 20px;
        font-weight: 600;
        font-size: 0.9rem;
    }}
    
    .match-score-high {{
        background-color: var(--success-green);
        color: white;
    }}
    
    .match-score-medium {{
        background-color: var(--warning-amber);
        color: white;
    }}
    
    .match-score-low {{
        background-color: var(--error-red);
        color: white;
    }}
    
    .expandable-details {{
        background-color: var(--bg-container);
        padding: 1.5rem;
        border-radius: 8px;
        margin-top: 1rem;
        border-left: 4px solid var(--primary-accent);
    }}
    
    .match-breakdown {{
        display: flex;
        gap: 2rem;
        margin: 1rem 0;
    }}
    
    .match-type {{
        flex: 1;
        padding: 1rem;
        background-color: var(--bg-main);
        border-radius: 8px;
    }}
    
    .match-type-label {{
        font-size: 0.85rem;
        color: var(--text-secondary);
        text-transform: uppercase;
        margin-bottom: 0.5rem;
    }}
    
    .match-type-value {{
        font-size: 1.5rem;
        font-weight: bold;
        color: var(--primary-accent);
    }}
    
    /* Sidebar styling */
    .sidebar-section {{
        margin-bottom: 2rem;
        padding-bottom: 1.5rem;
        border-bottom: 1px solid rgba(0, 0, 0, 0.1);
    }}
    
    [data-theme="dark"] .sidebar-section {{
        border-bottom: 1px solid rgba(255, 255, 255, 0.1);
    }}
    
    /* Table styling */
    .dataframe {{
        width: 100%;
        border-collapse: collapse;
    }}
    
    .dataframe th {{
        background-color: var(--bg-container);
        color: var(--text-primary);
        font-weight: 600;
        padding: 1rem;
        text-align: left;
        border-bottom: 2px solid var(--primary-accent);
    }}
    
    .dataframe td {{
        padding: 1rem;
        border-bottom: 1px solid rgba(0, 0, 0, 0.05);
    }}
    
    [data-theme="dark"] .dataframe td {{
        border-bottom: 1px solid rgba(255, 255, 255, 0.1);
    }}
    
    /* Missing Critical Skill column styling */
    .dataframe td:has-text("Missing Critical Skill"),
    .dataframe [data-column="Missing Critical Skill"] {{
        color: #FF6B6B !important;
        font-weight: 500;
    }}
    
    [data-theme="dark"] .dataframe td:has-text("Missing Critical Skill"),
    [data-theme="dark"] .dataframe [data-column="Missing Critical Skill"] {{
        color: #FF8C8C !important;
    }}
</style>
""", unsafe_allow_html=True)

# Auto-detect and apply theme based on system preferences
# This script runs immediately on page load to detect browser/system theme preference
st.markdown("""
<script>
(function() {
    // Function to update theme based on system preference
    function updateTheme() {
        const prefersDark = window.matchMedia('(prefers-color-scheme: dark)').matches;
        
        if (prefersDark) {
            // Apply dark theme
            document.documentElement.setAttribute('data-theme', 'dark');
            // Also apply to Streamlit app container if it exists
            const stApp = document.querySelector('.stApp') || document.querySelector('[data-testid="stApp"]');
            if (stApp) {
                stApp.setAttribute('data-theme', 'dark');
            }
            // Apply to body as well for comprehensive coverage
            document.body.setAttribute('data-theme', 'dark');
        } else {
            // Apply light theme (remove dark attribute)
            document.documentElement.removeAttribute('data-theme');
            const stApp = document.querySelector('.stApp') || document.querySelector('[data-testid="stApp"]');
            if (stApp) {
                stApp.removeAttribute('data-theme');
            }
            document.body.removeAttribute('data-theme');
        }
    }
    
    // Set initial theme immediately (before DOM is fully loaded to prevent flash)
    if (document.readyState === 'loading') {
        document.addEventListener('DOMContentLoaded', updateTheme);
    } else {
        updateTheme();
    }
    
    // Listen for changes in system theme preference
    const mediaQuery = window.matchMedia('(prefers-color-scheme: dark)');
    
    // Modern browsers support addEventListener
    if (mediaQuery.addEventListener) {
        mediaQuery.addEventListener('change', updateTheme);
    } else {
        // Fallback for older browsers (Safari < 14, etc.)
        mediaQuery.addListener(updateTheme);
    }
    
    // Also update theme when Streamlit app is ready (for dynamic content)
    if (window.parent !== window) {
        // Running in iframe (Streamlit Cloud)
        window.addEventListener('load', updateTheme);
    }
    
    // Use MutationObserver instead of setInterval for better performance
    // This watches for DOM changes and only checks theme when needed
    if (typeof MutationObserver !== 'undefined') {
        let lastThemeCheck = 0;
        const themeCheckThrottle = 2000; // Check at most once every 2 seconds
        
        const themeObserver = new MutationObserver(function() {
            const now = Date.now();
            if (now - lastThemeCheck < themeCheckThrottle) {
                return; // Throttle checks
            }
            lastThemeCheck = now;
            
            const currentPrefersDark = window.matchMedia('(prefers-color-scheme: dark)').matches;
            const hasDarkTheme = document.documentElement.hasAttribute('data-theme');
            
            if ((currentPrefersDark && !hasDarkTheme) || (!currentPrefersDark && hasDarkTheme)) {
                updateTheme();
            }
        });
        
        // Observe document body for changes (Streamlit dynamically loads content)
        themeObserver.observe(document.body, { 
            childList: true, 
            subtree: true,
            attributes: false // We don't need to watch attribute changes
        });
    }
})();
</script>
""", unsafe_allow_html=True)

# WebSocket connection stability and auto-reconnection handling
# This script handles WebSocket disconnections gracefully and auto-reconnects
st.markdown("""
<style>
    /* Connection status indicator styles */
    .ws-reconnecting-overlay {
        position: fixed;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: rgba(0, 0, 0, 0.7);
        display: flex;
        align-items: center;
        justify-content: center;
        z-index: 99999;
        opacity: 0;
        visibility: hidden;
        transition: opacity 0.3s, visibility 0.3s;
    }
    .ws-reconnecting-overlay.active {
        opacity: 1;
        visibility: visible;
    }
    .ws-reconnecting-content {
        background: white;
        padding: 30px 40px;
        border-radius: 12px;
        text-align: center;
        box-shadow: 0 10px 40px rgba(0,0,0,0.3);
    }
    [data-theme="dark"] .ws-reconnecting-content {
        background: #262626;
        color: #f4f4f4;
    }
    .ws-reconnecting-spinner {
        width: 40px;
        height: 40px;
        border: 4px solid #e0e0e0;
        border-top-color: #0F62FE;
        border-radius: 50%;
        animation: ws-spin 1s linear infinite;
        margin: 0 auto 15px;
    }
    @keyframes ws-spin {
        to { transform: rotate(360deg); }
    }
    .ws-reconnecting-text {
        font-size: 16px;
        color: #333;
        margin-bottom: 5px;
    }
    [data-theme="dark"] .ws-reconnecting-text {
        color: #f4f4f4;
    }
    .ws-reconnecting-subtext {
        font-size: 13px;
        color: #666;
    }
    [data-theme="dark"] .ws-reconnecting-subtext {
        color: #999;
    }
</style>
<div id="ws-reconnecting-overlay" class="ws-reconnecting-overlay">
    <div class="ws-reconnecting-content">
        <div class="ws-reconnecting-spinner"></div>
        <div class="ws-reconnecting-text">Reconnecting...</div>
        <div class="ws-reconnecting-subtext">Please wait while we restore your connection</div>
    </div>
</div>
<script>
(function() {
    // WebSocket connection monitor and auto-reconnection handler
    // Optimized for Streamlit Cloud deployment stability
    const overlay = document.getElementById('ws-reconnecting-overlay');
    let isReconnecting = false;
    let reconnectAttempts = 0;
    const maxReconnectAttempts = 8;  // Increased for better resilience
    let reconnectTimer = null;
    let connectionCheckTimer = null;
    let lastErrorTime = 0;
    const ERROR_DEBOUNCE_MS = 2000;  // Debounce rapid error events
    
    function showReconnectingOverlay() {
        // Debounce rapid reconnection triggers
        const now = Date.now();
        if (now - lastErrorTime < ERROR_DEBOUNCE_MS) {
            return;  // Skip if we just showed the overlay
        }
        lastErrorTime = now;
        
        if (overlay && !isReconnecting) {
            isReconnecting = true;
            overlay.classList.add('active');
        }
    }
    
    function hideReconnectingOverlay() {
        if (overlay) {
            isReconnecting = false;
            reconnectAttempts = 0;
            overlay.classList.remove('active');
        }
    }
    
    function attemptReconnect() {
        // Clear any existing timer
        if (reconnectTimer) {
            clearTimeout(reconnectTimer);
        }
        
        if (reconnectAttempts >= maxReconnectAttempts) {
            // Max attempts reached, reload the page
            console.log('CareerLens: Max reconnection attempts reached, reloading page...');
            // Use soft reload first (preserves some state)
            try {
                window.location.href = window.location.href;
            } catch (e) {
                window.location.reload();
            }
            return;
        }
        
        reconnectAttempts++;
        console.log('CareerLens: Reconnection attempt ' + reconnectAttempts + '/' + maxReconnectAttempts);
        
        // Exponential backoff with jitter: base * 2^attempt + random(0-500ms)
        const baseDelay = 1500;  // Start with 1.5s
        const delay = Math.min(baseDelay * Math.pow(1.5, reconnectAttempts - 1), 20000) + Math.random() * 500;
        
        reconnectTimer = setTimeout(function() {
            // Check if we're back online
            if (navigator.onLine) {
                // Check if Streamlit has already reconnected (look for active streams)
                if (activeConnections > 0) {
                    console.log('CareerLens: Connection restored via active WebSocket');
                    hideReconnectingOverlay();
                    return;
                }
                
                // Try to reload the Streamlit connection by triggering a rerun
                // This is safer than a full page reload
                try {
                    // Look for Streamlit's internal rerun mechanism
                    const stApp = window.parent.document || document;
                    const rerunButton = stApp.querySelector('[data-testid="stRerunButton"]');
                    if (rerunButton) {
                        rerunButton.click();
                        hideReconnectingOverlay();
                        return;
                    }
                } catch (e) {
                    console.log('CareerLens: Could not trigger Streamlit rerun');
                }
                
                // If we still can't reconnect after several attempts, reload
                if (reconnectAttempts >= 5) {
                    window.location.reload();
                } else {
                    attemptReconnect();
                }
            } else {
                // Still offline, keep trying
                attemptReconnect();
            }
        }, delay);
    }
    
    // Monitor network status
    window.addEventListener('offline', function() {
        console.log('CareerLens: Network connection lost');
        showReconnectingOverlay();
    });
    
    window.addEventListener('online', function() {
        console.log('CareerLens: Network connection restored');
        // Wait a moment for connection to stabilize
        setTimeout(function() {
            hideReconnectingOverlay();
            window.location.reload();
        }, 2000);
    });
    
    // Cleanup on page unload
    window.addEventListener('beforeunload', function() {
        if (reconnectTimer) {
            clearTimeout(reconnectTimer);
        }
        if (connectionCheckTimer) {
            clearInterval(connectionCheckTimer);
        }
    });
    
    console.log('CareerLens: Connection monitor initialized');
})();
</script>
""", unsafe_allow_html=True)

if 'search_history' not in st.session_state:
    st.session_state.search_history = []
if 'jobs_cache' not in st.session_state:
    st.session_state.jobs_cache = {}
if 'user_profile' not in st.session_state:
    st.session_state.user_profile = {}
if 'generated_resume' not in st.session_state:
    st.session_state.generated_resume = None
if 'selected_job' not in st.session_state:
    st.session_state.selected_job = None
if 'show_resume_generator' not in st.session_state:
    st.session_state.show_resume_generator = False
if 'resume_text' not in st.session_state:
    st.session_state.resume_text = None
if 'resume_embedding' not in st.session_state:
    st.session_state.resume_embedding = None
if 'matched_jobs' not in st.session_state:
    st.session_state.matched_jobs = []
if 'match_score' not in st.session_state:
    st.session_state.match_score = None
if 'missing_keywords' not in st.session_state:
    st.session_state.missing_keywords = None
if 'show_profile_editor' not in st.session_state:
    st.session_state.show_profile_editor = False
if 'use_auto_match' not in st.session_state:
    st.session_state.use_auto_match = False
if 'expanded_job_index' not in st.session_state:
    st.session_state.expanded_job_index = None
if 'industry_filter' not in st.session_state:
    st.session_state.industry_filter = None
if 'salary_min' not in st.session_state:
    st.session_state.salary_min = None
if 'salary_max' not in st.session_state:
    st.session_state.salary_max = None
if 'selected_job_index' not in st.session_state:
    st.session_state.selected_job_index = None
if 'dashboard_ready' not in st.session_state:
    st.session_state.dashboard_ready = False
if 'user_skills_embeddings_cache' not in st.session_state:
    st.session_state.user_skills_embeddings_cache = {}  # Cache for user skill embeddings
if 'skill_embeddings_cache' not in st.session_state:
    st.session_state.skill_embeddings_cache = {}  # General skill embedding cache

# Limit search history size to prevent memory bloat
MAX_SEARCH_HISTORY = 20
if len(st.session_state.search_history) > MAX_SEARCH_HISTORY:
    st.session_state.search_history = st.session_state.search_history[-MAX_SEARCH_HISTORY:]

# Run memory cleanup after session state is initialized
_cleanup_session_state()

# Flag to use fast string-based skill matching (default: True for speed)
# Set to False to use slower but more accurate semantic matching
USE_FAST_SKILL_MATCHING = os.getenv("USE_FAST_SKILL_MATCHING", "true").lower() in ("true", "1", "yes")

class APIMEmbeddingGenerator:
    def __init__(self, api_key, endpoint):
        self.api_key = api_key
        # Normalize endpoint: remove trailing slash
        endpoint = endpoint.rstrip('/')
        # Remove /openai if it's already in the endpoint (to avoid duplication)
        if endpoint.endswith('/openai'):
            endpoint = endpoint[:-7]  # Remove '/openai'
        self.endpoint = endpoint
        self.deployment = "text-embedding-3-small"
        self.api_version = "2024-02-01"
        self.url = f"{self.endpoint}/openai/deployments/{self.deployment}/embeddings?api-version={self.api_version}"
        self.headers = {"api-key": self.api_key, "Content-Type": "application/json"}
        self.encoding = tiktoken.get_encoding("cl100k_base")  # For token counting
    
    def get_embedding(self, text):
        """
        Generate embedding for a single text.
        
        Returns:
            tuple: (embedding, tokens_used) or (None, 0) on error
            - embedding: The embedding vector or None
            - tokens_used: Number of tokens used (from API response or estimated)
        """
        try:
            payload = {"input": text, "model": self.deployment}
            
            # Estimate tokens for fallback tracking
            estimated_tokens = len(self.encoding.encode(text))
            
            def make_request():
                return requests.post(self.url, headers=self.headers, json=payload, timeout=30)
            
            # api_call_with_retry handles rate limiting with exponential backoff
            response = api_call_with_retry(make_request, max_retries=3)
            
            if response and response.status_code == 200:
                result = response.json()
                embedding = result['data'][0]['embedding']
                
                # Return token usage from API response or fallback to estimate
                tokens_used = 0
                if 'usage' in result:
                    tokens_used = result['usage'].get('total_tokens', 0)
                else:
                    tokens_used = estimated_tokens
                
                return embedding, tokens_used
            else:
                # Error case - return None and 0 tokens
                return None, 0
        except Exception as e:
            st.error(f"Error generating embedding: {e}")
            return None, 0
    
    def get_embeddings_batch(self, texts, batch_size=None):
        """
        Generate embeddings for a batch of texts.
        
        Returns:
            tuple: (embeddings, total_tokens_used)
            - embeddings: List of embedding vectors
            - total_tokens_used: Total number of tokens used across all batches
        """
        if not texts:
            return [], 0
        
        effective_batch_size = batch_size or DEFAULT_EMBEDDING_BATCH_SIZE
        if effective_batch_size <= 0:
            effective_batch_size = DEFAULT_EMBEDDING_BATCH_SIZE
        
        embeddings = []
        total_tokens_used = 0
        progress_bar = st.progress(0)
        status_text = st.empty()
        total_batches = (len(texts) + effective_batch_size - 1) // effective_batch_size
        
        for i in range(0, len(texts), effective_batch_size):
            batch = texts[i:i + effective_batch_size]
            batch_num = i // effective_batch_size + 1
            progress = (i + len(batch)) / len(texts)
            progress_bar.progress(progress)
            status_text.text(f"🔄 Generating embeddings: {i + len(batch)}/{len(texts)} (batch {batch_num}/{total_batches})")
            
            # Add delay between batches to avoid rate limits (except for first batch)
            # Note: api_call_with_retry handles 429 errors with exponential backoff,
            # so this is just a gentle spacing between successful batches
            if i > 0 and EMBEDDING_BATCH_DELAY > 0:
                # Use chunked sleep to keep WebSocket alive during delays
                _chunked_sleep(EMBEDDING_BATCH_DELAY, f"Batch {batch_num}/{total_batches}")
            
            try:
                payload = {"input": batch, "model": self.deployment}
                
                # Estimate tokens for fallback tracking
                estimated_batch_tokens = sum(len(self.encoding.encode(text)) for text in batch)
                
                # Send keepalive before potentially long API call
                _websocket_keepalive(f"Processing batch {batch_num}/{total_batches}...")
                
                def make_request():
                    return requests.post(self.url, headers=self.headers, json=payload, timeout=25)
                
                # api_call_with_retry handles rate limiting with exponential backoff
                response = api_call_with_retry(make_request, max_retries=3)
                
                if response and response.status_code == 200:
                    data = response.json()
                    sorted_data = sorted(data['data'], key=lambda x: x['index'])
                    embeddings.extend([item['embedding'] for item in sorted_data])
                    
                    # Get token usage from API response or fallback to estimate
                    if 'usage' in data:
                        tokens_used = data['usage'].get('total_tokens', 0)
                    else:
                        tokens_used = estimated_batch_tokens
                    total_tokens_used += tokens_used
                elif response and response.status_code == 429:
                    # Rate limit hit - api_call_with_retry already handled retries with backoff
                    # If still 429 after retries, skip this batch to avoid further rate limit issues
                    st.warning(f"⚠️ Rate limit reached after retries. Skipping batch {batch_num}/{total_batches}. Consider reducing batch size or waiting before retrying.")
                    # Don't try individual calls - they'll likely hit the same rate limit
                    # This prevents cascading rate limit failures
                else:
                    # Other error - fallback to individual calls
                    st.warning(f"⚠️ Batch embedding failed (status {response.status_code if response else 'None'}), trying individual calls for batch {batch_num}...")
                    for text in batch:
                        emb, tokens = self.get_embedding(text)
                        if emb:
                            embeddings.append(emb)
                            total_tokens_used += tokens
            except Exception as e:
                st.warning(f"⚠️ Error processing batch {batch_num}, trying individual calls: {e}")
                for text in batch:
                    emb, tokens = self.get_embedding(text)
                    if emb:
                        embeddings.append(emb)
                        total_tokens_used += tokens
        
        progress_bar.empty()
        status_text.empty()
        return embeddings, total_tokens_used

class AzureOpenAITextGenerator:
    def __init__(self, api_key, endpoint, token_tracker=None):
        self.api_key = api_key
        # Normalize endpoint: remove trailing slash
        endpoint = endpoint.rstrip('/')
        # Remove /openai if it's already in the endpoint (to avoid duplication)
        if endpoint.endswith('/openai'):
            endpoint = endpoint[:-7]  # Remove '/openai'
        self.endpoint = endpoint
        self.deployment = "gpt-4o-mini"  # or your deployment name
        self.api_version = "2024-02-01"
        self.url = f"{self.endpoint}/openai/deployments/{self.deployment}/chat/completions?api-version={self.api_version}"
        self.headers = {"api-key": self.api_key, "Content-Type": "application/json"}
        self.token_tracker = token_tracker
        self.encoding = tiktoken.get_encoding("cl100k_base")  # For token counting
    
    def generate_resume(self, user_profile, job_posting, raw_resume_text=None):
        """Generate a tailored resume based on user profile and job posting using Context Sandwich approach.
        Returns structured JSON data instead of formatted text."""
        
        # System Instructions
        system_instructions = """You are an expert resume writer with expertise in ATS optimization and career coaching.
Your task is to create a tailored resume by analyzing the job description and adapting the user's profile.
Return ONLY valid JSON - no markdown, no additional text, no code blocks."""

        # Job Description Section
        job_description = f"""JOB POSTING TO MATCH:
Title: {job_posting.get('title', 'N/A')}
Company: {job_posting.get('company', 'N/A')}
Description: {job_posting.get('description', 'N/A')}
Required Skills: {', '.join(job_posting.get('skills', []))}"""

        # Structured Profile Section
        structured_profile = f"""STRUCTURED PROFILE:
Name: {user_profile.get('name', 'N/A')}
Email: {user_profile.get('email', 'N/A')}
Phone: {user_profile.get('phone', 'N/A')}
Location: {user_profile.get('location', 'N/A')}
LinkedIn: {user_profile.get('linkedin', 'N/A')}
Portfolio: {user_profile.get('portfolio', 'N/A')}
Summary: {user_profile.get('summary', 'N/A')}
Experience: {user_profile.get('experience', 'N/A')}
Education: {user_profile.get('education', 'N/A')}
Skills: {user_profile.get('skills', 'N/A')}
Certifications: {user_profile.get('certifications', 'N/A')}"""

        # Raw Original Resume (if available)
        raw_resume_section = ""
        if raw_resume_text:
            raw_resume_section = f"\n\nORIGINAL RESUME TEXT (for reference and context):\n{raw_resume_text[:3000]}"  # Limit to avoid token limits

        # Context Sandwich: System Instructions + Job Description + (Structured Profile + Raw Resume)
        prompt = f"""{system_instructions}

{job_description}

{structured_profile}{raw_resume_section}

INSTRUCTIONS:
1. Analyze the job posting requirements and identify key skills, technologies, and qualifications needed
2. Tailor the user's profile to match the job description by:
   - Rewriting the summary to emphasize relevant experience
   - Highlighting skills that match the job requirements
   - Rewriting experience bullet points to emphasize relevant achievements
   - Using keywords from the job description for ATS optimization
3. Focus on achievements and measurable results
4. Maintain accuracy - only use information from the provided profile

Return your response as a JSON object with this exact structure:
{{
  "header": {{
    "name": "Full Name",
    "title": "Professional Title (tailored to job)",
    "email": "email@example.com",
    "phone": "phone number",
    "location": "City, State/Country",
    "linkedin": "LinkedIn URL or empty string",
    "portfolio": "Portfolio URL or empty string"
  }},
  "summary": "2-3 sentence professional summary tailored to the job description, emphasizing relevant experience and skills",
  "skills_highlighted": ["Skill 1", "Skill 2", "Skill 3", ...],
  "experience": [
    {{
      "company": "Company Name",
      "title": "Job Title",
      "dates": "Date Range",
      "bullets": [
        "Rewritten bullet point emphasizing relevant achievement...",
        "Another tailored bullet point..."
      ]
    }}
  ],
  "education": "Education details formatted as text",
  "certifications": "Certifications, awards, or other achievements formatted as text"
}}

IMPORTANT: Return ONLY the JSON object, no markdown code blocks, no additional text."""
        
        try:
            payload = {
                "messages": [
                    {"role": "system", "content": system_instructions},
                    {"role": "user", "content": prompt}
                ],
                "max_tokens": 3000,
                "temperature": 0.7,
                "response_format": {"type": "json_object"}  # Force JSON output
            }
            
            # Send keepalive before potentially long LLM call
            _websocket_keepalive("Generating resume...")
            
            def make_request():
                return requests.post(self.url, headers=self.headers, json=payload, timeout=45)
            
            response = api_call_with_retry(make_request, max_retries=3)
            
            if response and response.status_code == 200:
                result = response.json()
                content = result['choices'][0]['message']['content']
                
                # Track token usage
                if self.token_tracker and 'usage' in result:
                    usage = result['usage']
                    prompt_tokens = usage.get('prompt_tokens', 0)
                    completion_tokens = usage.get('completion_tokens', 0)
                    self.token_tracker.add_completion_tokens(prompt_tokens, completion_tokens)
                
                # Parse JSON response
                try:
                    # Remove markdown code blocks if present
                    content = content.strip()
                    if content.startswith("```"):
                        lines = content.split('\n')
                        content = '\n'.join(lines[1:-1]) if lines[-1].startswith('```') else '\n'.join(lines[1:])
                    
                    resume_data = json.loads(content)
                    return resume_data
                except json.JSONDecodeError as e:
                    # Try to extract JSON from response
                    json_match = re.search(r'\{.*\}', content, re.DOTALL)
                    if json_match:
                        resume_data = json.loads(json_match.group())
                        return resume_data
                    else:
                        st.error(f"Could not parse JSON response: {e}")
                        return None
            else:
                if response:
                    error_detail = response.text[:200] if response.text else "No error details"
                    st.error(f"API Error: {response.status_code} - {error_detail}")
                return None
        except Exception as e:
            st.error(f"Error generating resume: {e}")
            return None
    
    def calculate_match_score(self, resume_content, job_description, embedding_generator):
        """Calculate match score between resume and job description, and identify missing keywords.
        Returns (None, None) if embeddings cannot be generated."""
        try:
            # Create embeddings for resume and job description
            resume_embedding, resume_tokens = embedding_generator.get_embedding(resume_content)
            job_embedding, job_tokens = embedding_generator.get_embedding(job_description)
            
            # Update token tracker if available
            token_tracker = get_token_tracker()
            if token_tracker:
                token_tracker.add_embedding_tokens(resume_tokens + job_tokens)
            
            if not resume_embedding or not job_embedding:
                return None, None
            
            # Calculate cosine similarity
            resume_emb = np.array(resume_embedding).reshape(1, -1)
            job_emb = np.array(job_embedding).reshape(1, -1)
            similarity = cosine_similarity(resume_emb, job_emb)[0][0]
            match_score = float(similarity)
            
            # Extract keywords from job description using AI
            # Use up to 8000 characters for keyword extraction (enough for most JDs while staying within API limits)
            job_desc_for_keywords = job_description[:8000] if len(job_description) > 8000 else job_description
            if len(job_description) > 8000:
                job_desc_for_keywords += "\n\n[Description truncated for keyword extraction - full description available for matching]"
            
            keyword_prompt = f"""Extract the most important technical skills, tools, technologies, and qualifications mentioned in this job description. 
Return ONLY a JSON object with a "keywords" array, no additional text.

Job Description:
{job_desc_for_keywords}

Return format: {{"keywords": ["keyword1", "keyword2", "keyword3", ...]}}"""
            
            payload = {
                "messages": [
                    {"role": "system", "content": "You are a keyword extraction expert. Extract only the most important technical and professional keywords. Return JSON with a 'keywords' array."},
                    {"role": "user", "content": keyword_prompt}
                ],
                "max_tokens": 500,
                "temperature": 0.3,
                "response_format": {"type": "json_object"}
            }
            
            def make_request():
                return requests.post(self.url, headers=self.headers, json=payload, timeout=30)
            
            response = api_call_with_retry(make_request, max_retries=2)
            
            missing_keywords = []
            if response and response.status_code == 200:
                try:
                    result = response.json()
                    content = result['choices'][0]['message']['content']
                    
                    # Track token usage
                    if self.token_tracker and 'usage' in result:
                        usage = result['usage']
                        prompt_tokens = usage.get('prompt_tokens', 0)
                        completion_tokens = usage.get('completion_tokens', 0)
                        self.token_tracker.add_completion_tokens(prompt_tokens, completion_tokens)
                    
                    # Try to parse keywords
                    keyword_data = json.loads(content)
                    job_keywords = keyword_data.get('keywords', [])
                    
                    # Check which keywords are missing from resume
                    resume_lower = resume_content.lower()
                    for keyword in job_keywords:
                        if isinstance(keyword, str) and keyword.lower() not in resume_lower:
                            missing_keywords.append(keyword)
                except Exception as e:
                    # If keyword extraction fails, continue without missing keywords
                    pass
            
            return match_score, missing_keywords[:10]  # Limit to top 10 missing keywords
            
        except Exception as e:
            st.warning(f"Could not calculate match score: {e}")
            return None, None
    
    def analyze_seniority_level(self, job_titles):
        """Analyze job titles to determine seniority level"""
        if not job_titles:
            return "Mid-Senior Level"
        
        titles_text = "\n".join([f"- {title}" for title in job_titles[:10]])
        prompt = f"""Analyze these job titles and determine the most common seniority level.
        
Job Titles:
{titles_text}

Return ONLY a JSON object with this structure:
{{
    "seniority": "Entry Level" | "Mid Level" | "Mid-Senior Level" | "Senior Level" | "Executive Level",
    "confidence": "high" | "medium" | "low"
}}

Choose the most appropriate seniority level based on the job titles."""
        
        try:
            payload = {
                "messages": [
                    {"role": "system", "content": "You are a career analyst. Analyze job titles and return only JSON."},
                    {"role": "user", "content": prompt}
                ],
                "max_tokens": 200,
                "temperature": 0.3,
                "response_format": {"type": "json_object"}
            }
            
            def make_request():
                return requests.post(self.url, headers=self.headers, json=payload, timeout=30)
            
            response = api_call_with_retry(make_request, max_retries=2)
            if response and response.status_code == 200:
                result = response.json()
                content = result['choices'][0]['message']['content']
                
                # Track token usage
                if self.token_tracker and 'usage' in result:
                    usage = result['usage']
                    prompt_tokens = usage.get('prompt_tokens', 0)
                    completion_tokens = usage.get('completion_tokens', 0)
                    self.token_tracker.add_completion_tokens(prompt_tokens, completion_tokens)
                
                data = json.loads(content)
                return data.get('seniority', 'Mid-Senior Level')
        except:
            pass
        
        # Fallback: simple keyword matching
        all_titles = " ".join(job_titles).lower()
        if any(word in all_titles for word in ['executive', 'director', 'vp', 'vice president', 'head of']):
            return "Executive Level"
        elif any(word in all_titles for word in ['senior', 'sr.', 'lead', 'principal']):
            return "Senior Level"
        elif any(word in all_titles for word in ['junior', 'jr.', 'entry', 'associate', 'graduate']):
            return "Entry Level"
        else:
            return "Mid-Senior Level"
    
    def recommend_accreditations(self, job_descriptions, user_skills):
        """Recommend accreditations based on job requirements"""
        if not job_descriptions:
            return "PMP or Scrum Master"
        
        # Combine job descriptions (limit to avoid token limits)
        combined_desc = "\n\n".join([desc[:1000] for desc in job_descriptions[:5]])
        user_skills_str = user_skills if user_skills else "Not specified"
        
        prompt = f"""Analyze these job descriptions and recommend the most valuable professional accreditation or certification for Hong Kong market.

Job Descriptions:
{combined_desc}

User's Current Skills: {user_skills_str}

Return ONLY a JSON object:
{{
    "accreditation": "Name of certification (e.g., PMP, HKICPA, AWS Certified)",
    "reason": "Brief reason why this certification is valuable"
}}

Focus on certifications that are:
1. Highly valued in Hong Kong market
2. Frequently mentioned in these job descriptions
3. Would unlock more opportunities for the user"""
        
        try:
            payload = {
                "messages": [
                    {"role": "system", "content": "You are a career advisor specializing in Hong Kong market. Return only JSON."},
                    {"role": "user", "content": prompt}
                ],
                "max_tokens": 300,
                "temperature": 0.5,
                "response_format": {"type": "json_object"}
            }
            
            def make_request():
                return requests.post(self.url, headers=self.headers, json=payload, timeout=30)
            
            response = api_call_with_retry(make_request, max_retries=2)
            if response and response.status_code == 200:
                result = response.json()
                content = result['choices'][0]['message']['content']
                
                # Track token usage
                if self.token_tracker and 'usage' in result:
                    usage = result['usage']
                    prompt_tokens = usage.get('prompt_tokens', 0)
                    completion_tokens = usage.get('completion_tokens', 0)
                    self.token_tracker.add_completion_tokens(prompt_tokens, completion_tokens)
                
                data = json.loads(content)
                return data.get('accreditation', 'PMP or Scrum Master')
        except:
            pass
        
        return "PMP or Scrum Master"
    
    def generate_recruiter_note(self, job, user_profile, semantic_score, skill_score):
        """Generate a personalized recruiter note"""
        job_title = job.get('title', '')
        job_desc = job.get('description', '')[:2000]  # Limit length
        user_summary = user_profile.get('summary', '')[:500]
        user_experience = user_profile.get('experience', '')[:500]
        
        prompt = f"""You are a professional recruiter in Hong Kong. Write a brief, actionable note about why this candidate is a good fit for this role.

Job Title: {job_title}
Job Description (excerpt): {job_desc}

Candidate Summary: {user_summary}
Candidate Experience (excerpt): {user_experience}

Match Scores:
- Semantic Match: {semantic_score:.0%}
- Skill Match: {skill_score:.0%}

Write a 2-3 sentence recruiter note that:
1. Highlights the strongest match points
2. Mentions any specific experience or skills that align well
3. Provides actionable feedback

Return ONLY the recruiter note text, no labels or formatting."""
        
        try:
            payload = {
                "messages": [
                    {"role": "system", "content": "You are a professional recruiter. Write concise, actionable notes."},
                    {"role": "user", "content": prompt}
                ],
                "max_tokens": 200,
                "temperature": 0.7
            }
            
            def make_request():
                return requests.post(self.url, headers=self.headers, json=payload, timeout=30)
            
            response = api_call_with_retry(make_request, max_retries=2)
            if response and response.status_code == 200:
                result = response.json()
                
                # Track token usage
                if self.token_tracker and 'usage' in result:
                    usage = result['usage']
                    prompt_tokens = usage.get('prompt_tokens', 0)
                    completion_tokens = usage.get('completion_tokens', 0)
                    self.token_tracker.add_completion_tokens(prompt_tokens, completion_tokens)
                
                return result['choices'][0]['message']['content'].strip()
        except:
            pass
        
        # Fallback note
        if semantic_score >= 0.7:
            return f"This role heavily emphasizes recent experience in {job.get('skills', ['relevant skills'])[0] if job.get('skills') else 'relevant skills'}, which is a strong point in your profile."
        else:
            return "Consider highlighting more relevant experience from your background to strengthen your application."

class RateLimiter:
    """Simple rate limiter that enforces requests per minute limit."""
    def __init__(self, max_requests_per_minute):
        self.max_requests_per_minute = max_requests_per_minute
        self.request_times = []
        self.lock = False  # Simple lock to prevent concurrent modifications
    
    def wait_if_needed(self):
        """Wait if we've exceeded the rate limit, otherwise record the request."""
        if self.max_requests_per_minute <= 0:
            return  # No rate limiting
        
        now = time.time()
        one_minute_ago = now - 60
        
        # Clean up old requests (older than 1 minute)
        self.request_times = [t for t in self.request_times if t > one_minute_ago]
        
        # If we're at the limit, wait until the oldest request is more than 1 minute old
        if len(self.request_times) >= self.max_requests_per_minute:
            oldest_request = min(self.request_times)
            wait_time = 60 - (now - oldest_request) + 1  # Add 1 second buffer
            if wait_time > 0:
                # Break long waits into smaller chunks to prevent WebSocket timeout
                # Streamlit Cloud's logstream WebSocket can disconnect after ~15-20s of inactivity
                # Using 2-second chunks ensures we update the UI frequently enough
                status_placeholder = st.empty()
                remaining = int(wait_time)
                while remaining > 0:
                    status_placeholder.info(f"⏳ Rate limiting: Waiting {remaining}s to stay under {self.max_requests_per_minute} requests/minute...")
                    chunk = min(2, remaining)  # Max 2 seconds at a time for WebSocket stability
                    time.sleep(chunk)
                    remaining -= chunk
                status_placeholder.empty()
                # Clean up again after waiting
                now = time.time()
                one_minute_ago = now - 60
                self.request_times = [t for t in self.request_times if t > one_minute_ago]
        
        # Record this request
        self.request_times.append(time.time())


class IndeedScraperAPI:
    """Job scraper using Indeed Scraper API via RapidAPI."""
    def __init__(self, api_key):
        self.api_key = api_key
        self.url = "https://indeed-scraper-api.p.rapidapi.com/api/job"
        self.headers = {
            'Content-Type': 'application/json',
            'x-rapidapi-host': 'indeed-scraper-api.p.rapidapi.com',
            'x-rapidapi-key': api_key
        }
        # Initialize rate limiter
        self.rate_limiter = RateLimiter(RAPIDAPI_MAX_REQUESTS_PER_MINUTE)
    
    def search_jobs(self, query, location="Hong Kong", max_rows=15, job_type="fulltime", country="hk"):
        payload = {
            "scraper": {
                "maxRows": max_rows,
                "query": query,
                "location": location,
                "jobType": job_type,
                "radius": "50",
                "sort": "relevance",
                "fromDays": "7",
                "country": country
            }
        }
        
        try:
            # Enforce rate limiting before making the request
            self.rate_limiter.wait_if_needed()
            
            # Send keepalive before potentially long API call
            _websocket_keepalive("Searching jobs...")
            
            def make_request():
                return requests.post(self.url, headers=self.headers, json=payload, timeout=45)
            
            response = api_call_with_retry(make_request, max_retries=3, initial_delay=3)
            
            if response and response.status_code == 201:
                data = response.json()
                jobs = []
                
                if 'returnvalue' in data and 'data' in data['returnvalue']:
                    job_list = data['returnvalue']['data']
                    
                    for job_data in job_list:
                        parsed_job = self._parse_job(job_data)
                        if parsed_job:
                            jobs.append(parsed_job)
                
                return jobs
            else:
                if response:
                    if response.status_code == 429:
                        st.error("🚫 Rate limit reached for Indeed API. Please wait a few minutes and try again.")
                    else:
                        error_detail = response.text[:200] if response.text else "No error details"
                        st.error(f"API Error: {response.status_code} - {error_detail}")
                return []
                
        except Exception as e:
            st.error(f"Error: {e}")
            return []
    
    def _parse_job(self, job_data):
        try:
            location_data = job_data.get('location', {})
            location = location_data.get('formattedAddressShort') or location_data.get('city', 'Hong Kong')
            
            job_types = job_data.get('jobType', [])
            job_type = ', '.join(job_types) if job_types else 'Full-time'
            
            benefits = job_data.get('benefits', [])
            attributes = job_data.get('attributes', [])
            
            # Get full description without truncation
            full_description = job_data.get('descriptionText', 'No description')
            # Store full description, but limit to 50000 chars to prevent memory issues
            description = full_description[:50000] if len(full_description) > 50000 else full_description
            
            return {
                'title': job_data.get('title', 'N/A'),
                'company': job_data.get('companyName', 'N/A'),
                'location': location,
                'description': description,
                'salary': 'Not specified',
                'job_type': job_type,
                'url': job_data.get('jobUrl', '#'),
                'posted_date': job_data.get('age', 'Recently'),
                'benefits': benefits[:5],
                'skills': attributes[:10],
                'company_rating': job_data.get('rating', {}).get('rating', 0),
                'is_remote': job_data.get('isRemote', False)
            }
        except:
            return None


class TokenUsageTracker:
    """Tracks token usage and costs for API calls."""
    def __init__(self):
        self.total_tokens = 0
        self.total_prompt_tokens = 0
        self.total_completion_tokens = 0
        self.total_embedding_tokens = 0
        self.cost_usd = 0.0
        # Pricing (as of 2024, adjust as needed)
        self.embedding_cost_per_1k = 0.00002  # $0.00002 per 1K tokens for text-embedding-3-small
        self.gpt4_mini_prompt_cost_per_1k = 0.00015  # $0.00015 per 1K tokens
        self.gpt4_mini_completion_cost_per_1k = 0.0006  # $0.0006 per 1K tokens
    
    def add_embedding_tokens(self, tokens):
        """Track embedding token usage."""
        self.total_embedding_tokens += tokens
        self.total_tokens += tokens
        self.cost_usd += (tokens / 1000) * self.embedding_cost_per_1k
    
    def add_completion_tokens(self, prompt_tokens, completion_tokens):
        """Track completion token usage."""
        self.total_prompt_tokens += prompt_tokens
        self.total_completion_tokens += completion_tokens
        self.total_tokens += prompt_tokens + completion_tokens
        self.cost_usd += (prompt_tokens / 1000) * self.gpt4_mini_prompt_cost_per_1k
        self.cost_usd += (completion_tokens / 1000) * self.gpt4_mini_completion_cost_per_1k
    
    def get_summary(self):
        """Get usage summary."""
        return {
            'total_tokens': self.total_tokens,
            'embedding_tokens': self.total_embedding_tokens,
            'prompt_tokens': self.total_prompt_tokens,
            'completion_tokens': self.total_completion_tokens,
            'estimated_cost_usd': self.cost_usd
        }
    
    def reset(self):
        """Reset counters."""
        self.total_tokens = 0
        self.total_prompt_tokens = 0
        self.total_completion_tokens = 0
        self.total_embedding_tokens = 0
        self.cost_usd = 0.0

def _is_streamlit_cloud():
    """Detect if running on Streamlit Cloud (ephemeral filesystem)."""
    # Streamlit Cloud sets STREAMLIT_SHARING_MODE or runs in specific environment
    return (
        os.environ.get('STREAMLIT_SHARING_MODE') is not None or
        os.environ.get('STREAMLIT_SERVER_PORT') is not None or
        os.path.exists('/mount/src') or  # Streamlit Cloud container path
        'streamlit.app' in os.environ.get('HOSTNAME', '')
    )


class SemanticJobSearch:
    def __init__(self, embedding_generator, use_persistent_store=True):
        self.embedding_gen = embedding_generator
        self.job_embeddings = []
        self.jobs = []
        self.chroma_client = None
        self.collection = None
        
        # On Streamlit Cloud, always use in-memory storage (ephemeral filesystem)
        if _is_streamlit_cloud():
            use_persistent_store = False
        
        # If ChromaDB is not available, skip vector store initialization
        if not CHROMADB_AVAILABLE:
            use_persistent_store = False
            self.use_persistent_store = False
            return  # Use simple list storage fallback
        
        self.use_persistent_store = use_persistent_store
        
        if use_persistent_store:
            try:
                # Initialize ChromaDB with persistent storage (local development only)
                chroma_db_path = os.path.join(os.getcwd(), ".chroma_db")
                os.makedirs(chroma_db_path, exist_ok=True)
                self.chroma_client = chromadb.PersistentClient(path=chroma_db_path)
                self.collection = self.chroma_client.get_or_create_collection(
                    name="job_embeddings",
                    metadata={"hnsw:space": "cosine"}
                )
            except Exception as e:
                st.warning(f"⚠️ Could not initialize persistent vector store: {e}. Using in-memory storage.")
                self.use_persistent_store = False
        
        # Use in-memory ChromaDB for Streamlit Cloud or fallback
        if not self.use_persistent_store and self.chroma_client is None:
            try:
                # EphemeralClient for in-memory storage (Streamlit Cloud compatible)
                self.chroma_client = chromadb.EphemeralClient()
                self.collection = self.chroma_client.get_or_create_collection(
                    name="job_embeddings",
                    metadata={"hnsw:space": "cosine"}
                )
            except Exception as e:
                # Final fallback: no ChromaDB, use simple list storage
                self.use_persistent_store = False
    
    def _get_job_hash(self, job):
        """Generate a hash for a job to use as ID."""
        job_str = f"{job.get('title', '')}_{job.get('company', '')}_{job.get('url', '')}"
        return hashlib.md5(job_str.encode()).hexdigest()
    
    def index_jobs(self, jobs, max_jobs_to_index=None):
        """Simplified job indexing: Check if job exists, if not, embed and store."""
        if not jobs:
            st.warning("⚠️ No jobs available to index.")
            self.jobs = []
            self.job_embeddings = []
            return
        
        effective_limit = max_jobs_to_index or min(len(jobs), DEFAULT_MAX_JOBS_TO_INDEX)
        effective_limit = max(1, min(effective_limit, len(jobs)))
        if effective_limit < len(jobs):
            st.info(f"⚙️ Indexing first {effective_limit} of {len(jobs)} jobs to reduce embedding API calls.")
        jobs_to_index = jobs[:effective_limit]
        self.jobs = jobs_to_index
        job_texts = [
            f"{job['title']} at {job['company']}. {job['description']} Skills: {', '.join(job['skills'][:5])}"
            for job in jobs_to_index
        ]
        
        st.info(f"📊 Indexing {len(jobs_to_index)} jobs...")
        
        # Simplified indexing: Check existing, generate missing, retrieve all
        if self.use_persistent_store and self.collection:
            try:
                # Generate hashes for all jobs
                job_hashes = [self._get_job_hash(job) for job in jobs_to_index]
                
                # Check which jobs already exist
                existing_data = self.collection.get(ids=job_hashes, include=['embeddings'])
                existing_ids = set(existing_data.get('ids', [])) if existing_data else set()
                
                # Find jobs that need embeddings
                jobs_to_embed = []
                indices_to_embed = []
                for idx, job_hash in enumerate(job_hashes):
                    if job_hash not in existing_ids:
                        jobs_to_embed.append(job_texts[idx])
                        indices_to_embed.append(idx)
                
                # Generate embeddings for new jobs only
                if jobs_to_embed:
                    st.info(f"🔄 Generating embeddings for {len(jobs_to_embed)} new jobs...")
                    new_embeddings, tokens_used = self.embedding_gen.get_embeddings_batch(jobs_to_embed)
                    
                    # Update token tracker
                    token_tracker = get_token_tracker()
                    if token_tracker:
                        token_tracker.add_embedding_tokens(tokens_used)
                    
                    # Store new embeddings
                    for idx, emb in zip(indices_to_embed, new_embeddings):
                        if emb:  # Only store if embedding was successfully generated
                            job_hash = job_hashes[idx]
                            self.collection.upsert(
                                ids=[job_hash],
                                embeddings=[emb],
                                documents=[job_texts[idx]],
                                metadatas=[{"job_index": idx}]
                            )
                
                # Retrieve all embeddings (existing + newly generated)
                retrieved = self.collection.get(ids=job_hashes, include=['embeddings'])
                if retrieved and 'embeddings' in retrieved and retrieved['embeddings'] is not None and len(retrieved['embeddings']) > 0:
                    # Map hashes to embeddings and maintain job order
                    hash_to_emb = {h: e for h, e in zip(retrieved['ids'], retrieved['embeddings'])}
                    self.job_embeddings = [hash_to_emb.get(h, None) for h in job_hashes]
                    # Filter out None values
                    self.job_embeddings = [e for e in self.job_embeddings if e is not None]
                    st.success(f"✅ Indexed {len(self.job_embeddings)} jobs (using persistent store)")
                else:
                    # Fallback: generate all if retrieval fails
                    self.job_embeddings, tokens_used = self.embedding_gen.get_embeddings_batch(job_texts)
                    token_tracker = get_token_tracker()
                    if token_tracker:
                        token_tracker.add_embedding_tokens(tokens_used)
                    st.success(f"✅ Indexed {len(self.job_embeddings)} jobs")
            except Exception as e:
                st.warning(f"⚠️ Error using persistent store: {e}. Generating new embeddings...")
                self.job_embeddings, tokens_used = self.embedding_gen.get_embeddings_batch(job_texts)
                token_tracker = get_token_tracker()
                if token_tracker:
                    token_tracker.add_embedding_tokens(tokens_used)
                self.use_persistent_store = False
                st.success(f"✅ Indexed {len(self.job_embeddings)} jobs")
        else:
            # Fallback to in-memory storage
            self.job_embeddings, tokens_used = self.embedding_gen.get_embeddings_batch(job_texts)
            token_tracker = get_token_tracker()
            if token_tracker:
                token_tracker.add_embedding_tokens(tokens_used)
            st.success(f"✅ Indexed {len(self.job_embeddings)} jobs")
    
    def search(self, query=None, top_k=10, resume_embedding=None):
        """Simplified search: Use pre-computed resume embedding if available, otherwise generate from query.
        
        Args:
            query: Optional query string (used if resume_embedding not provided)
            top_k: Number of top results to return
            resume_embedding: Pre-computed resume embedding (preferred, avoids API call)
        """
        if not self.job_embeddings:
            return []
        
        # Use pre-computed resume embedding if available (simplified - no caching needed)
        if resume_embedding is not None:
            query_embedding = resume_embedding
        elif query:
            # Fallback: generate embedding from query (for backward compatibility)
            query_embedding, tokens_used = self.embedding_gen.get_embedding(query)
            # Update token tracker
            token_tracker = get_token_tracker()
            if token_tracker:
                token_tracker.add_embedding_tokens(tokens_used)
            if not query_embedding:
                return []
        else:
            return []
        
        # Perform semantic similarity search
        query_emb = np.array(query_embedding).reshape(1, -1)
        job_embs = np.array(self.job_embeddings)
        
        similarities = cosine_similarity(query_emb, job_embs)[0]
        top_indices = np.argsort(similarities)[::-1][:top_k]
        
        results = []
        for idx in top_indices:
            results.append({
                'job': self.jobs[idx],
                'similarity_score': float(similarities[idx]),
                'rank': len(results) + 1
            })
        
        return results
    
    def calculate_skill_match(self, user_skills, job_skills):
        """Calculate skill-based match score.
        
        By default uses fast string-based matching (USE_FAST_SKILL_MATCHING=true).
        Set USE_FAST_SKILL_MATCHING=false for slower semantic matching with embeddings.
        """
        if not user_skills or not job_skills:
            return 0.0, []
        
        # Parse skills into lists
        user_skills_list = [s.strip() for s in str(user_skills).split(',') if s.strip()]
        job_skills_list = [s.strip() for s in job_skills if isinstance(s, str) and s.strip()]
        
        if not user_skills_list or not job_skills_list:
            return 0.0, []
        
        # Use fast string-based matching by default (no API calls, instant results)
        if USE_FAST_SKILL_MATCHING:
            return self._calculate_skill_match_string_based(user_skills_list, job_skills_list)
        
        # Semantic matching with embeddings (slower but more accurate)
        try:
            # Check cache for user skills embeddings
            user_skills_key = ",".join(sorted(user_skills_list))
            if user_skills_key in st.session_state.user_skills_embeddings_cache:
                user_skill_embeddings = st.session_state.user_skills_embeddings_cache[user_skills_key]
                user_tokens = 0  # No API call needed
            else:
                # Generate embeddings for user skills (use smaller batch size)
                user_skill_embeddings, user_tokens = self.embedding_gen.get_embeddings_batch(user_skills_list, batch_size=10)
                if user_skill_embeddings:
                    st.session_state.user_skills_embeddings_cache[user_skills_key] = user_skill_embeddings
            
            # Check cache for job skills embeddings
            job_skills_key = ",".join(sorted(job_skills_list))
            if job_skills_key in st.session_state.skill_embeddings_cache:
                job_skill_embeddings = st.session_state.skill_embeddings_cache[job_skills_key]
                job_tokens = 0  # No API call needed
            else:
                job_skill_embeddings, job_tokens = self.embedding_gen.get_embeddings_batch(job_skills_list, batch_size=10)
                if job_skill_embeddings:
                    st.session_state.skill_embeddings_cache[job_skills_key] = job_skill_embeddings
            
            # Update token tracker if available
            if user_tokens > 0 or job_tokens > 0:
                token_tracker = get_token_tracker()
                if token_tracker:
                    token_tracker.add_embedding_tokens(user_tokens + job_tokens)
            
            if not user_skill_embeddings or not job_skill_embeddings:
                # Fallback to string matching if embeddings fail
                return self._calculate_skill_match_string_based(user_skills_list, job_skills_list)
            
            # Convert to numpy arrays
            user_embs = np.array(user_skill_embeddings)
            job_embs = np.array(job_skill_embeddings)
            
            # Calculate similarity matrix (cosine similarity between all user skills and job skills)
            similarity_matrix = cosine_similarity(job_embs, user_embs)
            
            # Find best matches: for each job skill, find the best matching user skill
            # Use a threshold of 0.7 for semantic similarity (adjustable)
            similarity_threshold = 0.7
            matched_skills = []
            matched_indices = set()
            
            for job_idx, job_skill in enumerate(job_skills_list):
                best_match_idx = np.argmax(similarity_matrix[job_idx])
                best_similarity = similarity_matrix[job_idx][best_match_idx]
                
                if best_similarity >= similarity_threshold and best_match_idx not in matched_indices:
                    matched_skills.append(job_skill)
                    matched_indices.add(best_match_idx)
            
            # Calculate match percentage
            match_score = len(matched_skills) / len(job_skills_list) if job_skills_list else 0.0
            # Find missing skills (job skills that weren't matched)
            missing_skills = [js for js in job_skills_list if js not in matched_skills]
            
            return min(match_score, 1.0), missing_skills[:5]
            
        except Exception as e:
            # Fallback to string-based matching if semantic matching fails
            return self._calculate_skill_match_string_based(user_skills_list, job_skills_list)
    
    def _calculate_skill_match_string_based(self, user_skills_list, job_skills_list):
        """Fallback string-based skill matching"""
        user_skills_lower = [s.lower() for s in user_skills_list]
        job_skills_lower = [s.lower() for s in job_skills_list]
        
        matched_skills = []
        for job_skill in job_skills_lower:
            for user_skill in user_skills_lower:
                if job_skill in user_skill or user_skill in job_skill:
                    matched_skills.append(job_skill)
                    break
        
        match_score = len(matched_skills) / len(job_skills_lower) if job_skills_lower else 0.0
        missing_skills = [job_skills_list[i] for i, js in enumerate(job_skills_lower) if js not in matched_skills]
        
        return min(match_score, 1.0), missing_skills[:5]

def is_cache_valid(cache_entry):
    """Check if cache entry is still valid (not expired)."""
    if not cache_entry or not isinstance(cache_entry, dict):
        return False
    
    expires_at = cache_entry.get('expires_at')
    if expires_at is None:
        # Legacy cache without TTL - consider invalid for safety
        return False
    
    # Handle both datetime objects and string timestamps
    if isinstance(expires_at, str):
        try:
            expires_at = datetime.fromisoformat(expires_at)
        except ValueError:
            try:
                expires_at = datetime.strptime(expires_at, "%Y-%m-%d %H:%M:%S")
            except Exception:
                return False
    
    return datetime.now() < expires_at


def _build_jobs_cache_key(query, location, max_rows, job_type, country):
    """Create a unique cache key for job searches."""
    normalized_query = (query or "").strip().lower()
    return "|".join([
        normalized_query,
        (location or "").strip().lower(),
        str(max_rows),
        (job_type or "").strip().lower(),
        (country or "").strip().lower()
    ])


def _ensure_jobs_cache_structure():
    """Ensure jobs_cache is always a dict keyed by cache keys (handles legacy formats)."""
    if 'jobs_cache' not in st.session_state or not isinstance(st.session_state.jobs_cache, dict):
        st.session_state.jobs_cache = {}
        return
    cache = st.session_state.jobs_cache
    if cache and 'jobs' in cache and isinstance(cache['jobs'], list):
        cache_key = cache.get('cache_key') or _build_jobs_cache_key(
            cache.get('query', ''),
            cache.get('location', 'Hong Kong'),
            cache.get('count', len(cache.get('jobs', []))),
            cache.get('job_type', 'fulltime'),
            cache.get('country', 'hk')
        )
        st.session_state.jobs_cache = {cache_key: {**cache, 'cache_key': cache_key}}


def _get_cached_jobs(query, location, max_rows, job_type, country):
    """Return cached jobs for a given search signature if valid."""
    _ensure_jobs_cache_structure()
    cache_key = _build_jobs_cache_key(query, location, max_rows, job_type, country)
    cache_entry = st.session_state.jobs_cache.get(cache_key)
    if not cache_entry:
        return None
    if not is_cache_valid(cache_entry):
        st.session_state.jobs_cache.pop(cache_key, None)
        return None
    return cache_entry


def _store_jobs_in_cache(query, location, max_rows, job_type, country, jobs, cache_ttl_hours=168):
    """Persist job results in cache with TTL metadata."""
    _ensure_jobs_cache_structure()
    cache_key = _build_jobs_cache_key(query, location, max_rows, job_type, country)
    now = datetime.now()
    expires_at = now + timedelta(hours=cache_ttl_hours)
    st.session_state.jobs_cache[cache_key] = {
        'jobs': jobs,
        'count': len(jobs),
        'timestamp': now.isoformat(),
        'query': query,
        'location': location,
        'job_type': job_type,
        'country': country,
        'cache_key': cache_key,
        'expires_at': expires_at.isoformat()
    }
    return st.session_state.jobs_cache[cache_key]


def fetch_jobs_with_cache(scraper, query, location="Hong Kong", max_rows=25, job_type="fulltime",
                          country="hk", cache_ttl_hours=168, force_refresh=False):
    """
    Fetch jobs with session-level caching to avoid RapidAPI rate limits.
    Set force_refresh=True to bypass cache for a particular query.
    IMPORTANT: force_refresh should ONLY be True when user explicitly clicks a "Refresh Jobs" button,
    never implicitly based on timestamps or status checks.
    """
    if scraper is None:
        return []
    _ensure_jobs_cache_structure()
    cache_key = _build_jobs_cache_key(query, location, max_rows, job_type, country)
    if force_refresh:
        if cache_key in st.session_state.jobs_cache:
            st.caption("🔁 Forcing a fresh job search (cache bypassed)")
        st.session_state.jobs_cache.pop(cache_key, None)
    else:
        cache_entry = _get_cached_jobs(query, location, max_rows, job_type, country)
        if cache_entry:
            timestamp = cache_entry.get('timestamp')
            expires_at = cache_entry.get('expires_at')
            expires_in_minutes = None
            if isinstance(expires_at, str):
                try:
                    expires_dt = datetime.fromisoformat(expires_at)
                    expires_in_minutes = max(0, int((expires_dt - datetime.now()).total_seconds() // 60))
                except ValueError:
                    pass
            if timestamp and isinstance(timestamp, str):
                try:
                    ts_dt = datetime.fromisoformat(timestamp)
                    human_ts = ts_dt.strftime("%b %d %H:%M")
                except ValueError:
                    human_ts = timestamp
            else:
                human_ts = "earlier"
            remaining_text = f" (~{expires_in_minutes} min left)" if expires_in_minutes is not None else ""
            st.caption(f"♻️ Using cached job results from {human_ts}{remaining_text}")
            return cache_entry.get('jobs', [])
    jobs = scraper.search_jobs(query, location, max_rows, job_type, country)
    if jobs:
        _store_jobs_in_cache(query, location, max_rows, job_type, country, jobs, cache_ttl_hours)
    return jobs

def get_token_tracker():
    """Get or create token usage tracker."""
    if 'token_tracker' not in st.session_state:
        st.session_state.token_tracker = TokenUsageTracker()
    return st.session_state.token_tracker


@st.cache_resource(show_spinner=False)
def _create_embedding_generator_resource(api_key, endpoint):
    return APIMEmbeddingGenerator(api_key, endpoint)


@st.cache_resource(show_spinner=False)
def _create_text_generator_resource(api_key, endpoint):
    return AzureOpenAITextGenerator(api_key, endpoint)

def get_embedding_generator():
    """Get cached embedding generator instance.
    
    Uses @st.cache_resource to ensure the API client is created only once
    across the entire app lifespan, preventing unnecessary API calls.
    """
    try:
        # Use secrets instead of hardcoded values
        AZURE_OPENAI_API_KEY = st.secrets.get("AZURE_OPENAI_API_KEY")
        AZURE_OPENAI_ENDPOINT = st.secrets.get("AZURE_OPENAI_ENDPOINT")
        
        if not AZURE_OPENAI_API_KEY or not AZURE_OPENAI_ENDPOINT:
            st.error("⚠️ Azure OpenAI credentials are missing. Please configure AZURE_OPENAI_API_KEY and AZURE_OPENAI_ENDPOINT in your Streamlit secrets.")
            return None
        
        # Get cached generator instance (created only once via @st.cache_resource)
        generator = _create_embedding_generator_resource(AZURE_OPENAI_API_KEY, AZURE_OPENAI_ENDPOINT)
        
        # Note: Token tracking is now handled by callers, not by mutating the cached generator
        # This prevents issues with Streamlit's caching mechanism
        
        return generator
    except KeyError as e:
        st.error(f"⚠️ Missing required secret: {e}. Please configure your Streamlit secrets.")
        return None
    except Exception as e:
        st.error(f"⚠️ Error initializing embedding generator: {e}")
        return None

def generate_and_store_resume_embedding(resume_text, user_profile=None):
    """Generate embedding for resume and store in session state.
    
    This is called once when resume is uploaded/updated, so we can reuse
    the embedding for all subsequent searches without regenerating it.
    """
    if not resume_text:
        st.session_state.resume_embedding = None
        return None
    
    # Build resume query text (same logic as before, but we'll store the embedding)
    if user_profile:
        profile_data = f"{user_profile.get('summary', '')} {user_profile.get('experience', '')} {user_profile.get('skills', '')}"
        resume_query = f"{resume_text} {profile_data}"
    else:
        resume_query = resume_text
    
    # Generate embedding
    embedding_gen = get_embedding_generator()
    if not embedding_gen:
        return None
    
    embedding, tokens_used = embedding_gen.get_embedding(resume_query)
    # Update token tracker
    token_tracker = get_token_tracker()
    if token_tracker:
        token_tracker.add_embedding_tokens(tokens_used)
    
    if embedding:
        st.session_state.resume_embedding = embedding
        return embedding
    
    return None

def get_job_scraper():
    """Get Indeed job scraper.
    
    Uses RAPIDAPI_KEY to fetch jobs from Indeed Scraper API.
    """
    if 'job_scraper' not in st.session_state:
        RAPIDAPI_KEY = st.secrets.get("RAPIDAPI_KEY", "")
        if not RAPIDAPI_KEY:
            st.error("⚠️ RAPIDAPI_KEY is required in secrets. Please configure it in your .streamlit/secrets.toml")
            return None
        
        st.session_state.job_scraper = IndeedScraperAPI(RAPIDAPI_KEY)
    
    return st.session_state.job_scraper

def get_text_generator():
    """Get cached text generator instance.
    
    Uses @st.cache_resource to ensure the API client is created only once
    across the entire app lifespan, preventing unnecessary API calls.
    """
    try:
        AZURE_OPENAI_API_KEY = st.secrets.get("AZURE_OPENAI_API_KEY")
        AZURE_OPENAI_ENDPOINT = st.secrets.get("AZURE_OPENAI_ENDPOINT")
        
        if not AZURE_OPENAI_API_KEY or not AZURE_OPENAI_ENDPOINT:
            st.error("⚠️ Azure OpenAI credentials are missing. Please configure AZURE_OPENAI_API_KEY and AZURE_OPENAI_ENDPOINT in your Streamlit secrets.")
            return None
        
        # Get cached generator instance (created only once via @st.cache_resource)
        generator = _create_text_generator_resource(AZURE_OPENAI_API_KEY, AZURE_OPENAI_ENDPOINT)
        
        # Update token tracker reference (token_tracker itself is in session_state)
        generator.token_tracker = get_token_tracker()
        
        return generator
    except KeyError as e:
        st.error(f"⚠️ Missing required secret: {e}. Please configure your Streamlit secrets.")
        return None
    except Exception as e:
        st.error(f"⚠️ Error initializing text generator: {e}")
        return None

def extract_salary_from_text(text):
    """Extract salary information from job description text using LLM"""
    if not text:
        return None, None
    
    # Limit text length to avoid token limits (check first 3000 chars which usually contains salary info)
    text_for_extraction = text[:3000] if len(text) > 3000 else text
    
    try:
        text_gen = get_text_generator()
        if text_gen is None:
            return None, None
        
        prompt = f"""Extract salary information from this job description text. 
Look for salary ranges, amounts, and compensation details. Normalize everything to monthly HKD (Hong Kong Dollars).

JOB DESCRIPTION TEXT:
{text_for_extraction}

Extract and return salary information as JSON with this structure:
{{
    "min_salary_hkd_monthly": <number or null>,
    "max_salary_hkd_monthly": <number or null>,
    "found": true/false,
    "raw_text": "the exact salary text found in the description"
}}

Rules:
- Convert all amounts to monthly HKD (multiply annual by 12, weekly by 4.33, daily by 22)
- If only one amount is found, set both min and max to that value
- If a range is found (e.g., "60k-80k"), extract both min and max
- Handle formats like "competitive", "based on experience", "around 60k-80k annually" by extracting the numeric range
- If no salary is found, set "found": false and return null for min/max
- Always return valid JSON, no additional text"""

        payload = {
            "messages": [
                {"role": "system", "content": "You are a salary extraction expert. Extract salary information and normalize to monthly HKD. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            "max_tokens": 300,
            "temperature": 0.1,
            "response_format": {"type": "json_object"}
        }
        
        def make_request():
            return requests.post(
                text_gen.url,
                headers=text_gen.headers,
                json=payload,
                timeout=30
            )
        
        response = api_call_with_retry(make_request, max_retries=2)
        
        if response and response.status_code == 200:
            result = response.json()
            content = result['choices'][0]['message']['content']
            
            # Track token usage
            if text_gen.token_tracker and 'usage' in result:
                usage = result['usage']
                prompt_tokens = usage.get('prompt_tokens', 0)
                completion_tokens = usage.get('completion_tokens', 0)
                text_gen.token_tracker.add_completion_tokens(prompt_tokens, completion_tokens)
            
            # Parse JSON response
            try:
                salary_data = json.loads(content)
                if salary_data.get('found', False):
                    min_sal = salary_data.get('min_salary_hkd_monthly')
                    max_sal = salary_data.get('max_salary_hkd_monthly')
                    if min_sal is not None and max_sal is not None:
                        return int(min_sal), int(max_sal)
                    elif min_sal is not None:
                        return int(min_sal), int(min_sal * 1.2)  # Estimate range
            except (json.JSONDecodeError, ValueError, TypeError) as e:
                # Fallback to regex if LLM parsing fails
                pass
        
        # Fallback to regex-based extraction if LLM fails
        return extract_salary_from_text_regex(text)
        
    except Exception as e:
        # Fallback to regex if LLM extraction fails
        return extract_salary_from_text_regex(text)

def extract_salary_from_text_regex(text):
    """Fallback regex-based salary extraction"""
    if not text:
        return None, None
    
    import re
    patterns = [
        r'HKD\s*\$?\s*(\d{1,3}(?:,\d{3})*(?:k|K)?)\s*[-–—]\s*\$?\s*(\d{1,3}(?:,\d{3})*(?:k|K)?)',
        r'(\d{1,3}(?:,\d{3})*(?:k|K)?)\s*[-–—]\s*(\d{1,3}(?:,\d{3})*(?:k|K)?)\s*HKD',
        r'HKD\s*\$?\s*(\d{1,3}(?:,\d{3})*(?:k|K)?)\s*(?:per month|/month|/mth|monthly)',
        r'(\d{1,3}(?:,\d{3})*(?:k|K)?)\s*HKD\s*(?:per month|/month|/mth|monthly)',
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            match = matches[0]
            if isinstance(match, tuple) and len(match) == 2:
                min_sal = match[0].replace(',', '').replace('k', '000').replace('K', '000')
                max_sal = match[1].replace(',', '').replace('k', '000').replace('K', '000')
                try:
                    min_val = int(min_sal)
                    max_val = int(max_sal)
                    return min_val, max_val
                except:
                    pass
            elif isinstance(match, tuple) and len(match) == 1:
                sal = match[0].replace(',', '').replace('k', '000').replace('K', '000')
                try:
                    sal_val = int(sal)
                    return sal_val, sal_val * 1.2  # Estimate range
                except:
                    pass
    
    return None, None

def calculate_salary_band(matched_jobs):
    """Calculate estimated salary band from matched jobs"""
    salaries = []
    
    for result in matched_jobs:
        job = result['job']
        # Try to extract from salary field
        salary_str = job.get('salary', '')
        if salary_str and salary_str != 'Not specified':
            min_sal, max_sal = extract_salary_from_text(salary_str)
            if min_sal and max_sal:
                salaries.append((min_sal, max_sal))
        
        # Try to extract from description
        description = job.get('description', '')
        if description:
            min_sal, max_sal = extract_salary_from_text(description[:5000])  # Check first 5000 chars
            if min_sal and max_sal:
                salaries.append((min_sal, max_sal))
    
    if not salaries:
        # Default estimate based on Hong Kong market
        return 45000, 55000
    
    # Calculate average min and max
    avg_min = int(np.mean([s[0] for s in salaries]))
    avg_max = int(np.mean([s[1] for s in salaries]))
    
    return avg_min, avg_max

def filter_jobs_by_domains(jobs, target_domains):
    """Filter jobs by target domains"""
    if not target_domains:
        return jobs
    
    filtered = []
    domain_keywords = {
        'FinTech': ['fintech', 'financial technology', 'blockchain', 'crypto', 'cryptocurrency', 'payment', 'banking technology', 'digital banking', 'wealthtech', 'insurtech'],
        'ESG & Sustainability': ['esg', 'sustainability', 'environmental', 'green', 'carbon', 'climate', 'renewable', 'sustainable'],
        'Data Analytics': ['data analytics', 'data analysis', 'business intelligence', 'bi', 'data science', 'data engineer', 'analytics', 'big data'],
        'Digital Transformation': ['digital transformation', 'digitalization', 'digital strategy', 'innovation', 'digital', 'transformation'],
        'Investment Banking': ['investment banking', 'ib', 'm&a', 'mergers', 'acquisitions', 'capital markets', 'equity research', 'corporate finance'],
        'Consulting': ['consulting', 'consultant', 'advisory', 'strategy consulting', 'management consulting'],
        'Technology': ['software', 'technology', 'tech', 'engineering', 'developer', 'programming', 'it', 'information technology', 'software engineer'],
        'Healthcare': ['healthcare', 'medical', 'health', 'hospital', 'clinical', 'pharmaceutical', 'biotech'],
        'Education': ['education', 'teaching', 'academic', 'university', 'school', 'e-learning', 'edtech'],
        'Real Estate': ['real estate', 'property', 'realty', 'property management', 'real estate development'],
        'Retail & E-commerce': ['retail', 'e-commerce', 'ecommerce', 'online retail', 'retail management'],
        'Marketing & Advertising': ['marketing', 'advertising', 'brand', 'digital marketing', 'social media marketing'],
        'Legal': ['legal', 'law', 'attorney', 'lawyer', 'compliance', 'regulatory'],
        'Human Resources': ['human resources', 'hr', 'recruitment', 'talent acquisition', 'people operations'],
        'Operations': ['operations', 'operations management', 'supply chain', 'logistics', 'procurement']
    }
    
    for job in jobs:
        # Check full description, not just first 2000 chars, and include company name
        title_lower = job.get('title', '').lower()
        desc_lower = job.get('description', '').lower()  # Check full description
        company_lower = job.get('company', '').lower()
        combined = f"{title_lower} {desc_lower} {company_lower}"
        
        for domain in target_domains:
            keywords = domain_keywords.get(domain, [domain.lower()])
            if any(keyword.lower() in combined for keyword in keywords):
                filtered.append(job)
                break  # Job matched a domain, no need to check other domains
    
    return filtered if filtered else jobs  # Return all if no matches

def filter_jobs_by_salary(jobs, min_salary):
    """Filter jobs by minimum salary expectation"""
    if not min_salary or min_salary <= 0:
        return jobs
    
    filtered = []
    jobs_without_salary = []
    
    for job in jobs:
        salary_str = job.get('salary', '')
        description = job.get('description', '')
        
        # Try to extract salary from salary field first
        min_sal, max_sal = extract_salary_from_text(salary_str)
        
        # If not found, try full description (not just first 5000 chars)
        if not min_sal:
            min_sal, max_sal = extract_salary_from_text(description)  # Check full description
        
        # If we found a salary
        if min_sal:
            # Include if minimum salary meets requirement OR if salary range overlaps with requirement
            # (e.g., if job is 50k-70k and user wants 60k, include it because 70k >= 60k)
            if min_sal >= min_salary or (max_sal and max_sal >= min_salary):
                filtered.append(job)
            # else: job salary is too low, exclude it
        else:
            # No salary info found - store separately to add at end if no matches
            jobs_without_salary.append(job)
    
    # If we have matches, return only matches
    # If no matches but we have jobs without salary, include those (user can decide)
    if filtered:
        return filtered
    elif jobs_without_salary:
        # No jobs met salary requirement, but return jobs without salary info
        # so user knows there are jobs available (they just don't have salary listed)
        return jobs_without_salary
    else:
        # No jobs at all
        return []

def display_job_card(result, index):
    job = result['job']
    score = result.get('similarity_score', 0.0)
    
    remote_badge = "🏠 Remote" if job['is_remote'] else ""
    rating = job['company_rating']
    stars = "⭐" * int(rating) if rating > 0 else ""
    
    st.markdown(f"""
    <div class="job-card">
        <div style="display: flex; justify-content: space-between; align-items: start; margin-bottom: 1rem;">
            <div style="flex-grow: 1;">
                <h3 style="margin: 0; color: var(--primary-accent);">#{index} {job['title']}</h3>
                <p style="margin: 0.5rem 0; color: var(--text-secondary); font-size: 0.95rem;">
                    🏢 <strong>{job['company']}</strong> {stars} • 📍 {job['location']} {remote_badge}
                </p>
            </div>
            <div class="match-score">
                {score:.1%} Match
            </div>
        </div>
        <div style="display: flex; gap: 1rem; flex-wrap: wrap; margin-bottom: 0.5rem; color: var(--text-secondary);">
            <span>⏰ {job['job_type']}</span>
            <span>💰 {job['salary']}</span>
            <span>📅 {job['posted_date']}</span>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        if job['benefits']:
            st.write("**Benefits:**")
            for benefit in job['benefits']:
                st.markdown(f'<span class="tag">✓ {benefit}</span>', unsafe_allow_html=True)
    
    with col2:
        if job['skills']:
            st.write("**Skills:**")
            skills_text = " ".join([f'<span class="tag">{skill}</span>' for skill in job['skills'][:8]])
            st.markdown(skills_text, unsafe_allow_html=True)
    
    st.write("")
    
    col1, col2 = st.columns([4, 1])
    
    with col1:
        with st.expander("📝 View Full Description", expanded=False):
            # Display full description with proper formatting
            description_text = job['description']
            if len(description_text) > 10000:
                st.info(f"📄 Full description ({len(description_text):,} characters)")
                # Use text area for very long descriptions to allow scrolling
                st.text_area(
                    "Job Description",
                    value=description_text,
                    height=400,
                    key=f"desc_{index}",
                    label_visibility="collapsed"
                )
            else:
                st.write(description_text)
    
    with col2:
        col2a, col2b = st.columns(2)
        with col2a:
            if job['url'] != '#':
                st.link_button("Apply →", job['url'], use_container_width=True)
        with col2b:
            if st.button("📄 Resume", key=f"resume_{index}", use_container_width=True, type="primary"):
                st.session_state.selected_job = job
                st.session_state.show_resume_generator = True
                st.rerun()

def extract_text_from_resume(uploaded_file):
    """Extract text from uploaded resume file (PDF, DOCX, or TXT)"""
    try:
        file_type = uploaded_file.name.split('.')[-1].lower()
        
        if file_type == 'pdf':
            # Check if PyPDF2 is available
            if not PYPDF2_AVAILABLE:
                st.error("PDF processing is not available. Please upload a TXT file instead.")
                return None
            # Extract text from PDF
            uploaded_file.seek(0)  # Reset file pointer
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        
        elif file_type == 'docx':
            # Check if python-docx is available
            if not DOCX_AVAILABLE:
                st.error("DOCX processing is not available. Please upload a TXT or PDF file instead.")
                return None
            # Extract text from DOCX
            uploaded_file.seek(0)  # Reset file pointer
            doc = Document(uploaded_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        
        elif file_type == 'txt':
            # Read text file
            uploaded_file.seek(0)  # Reset file pointer
            text = str(uploaded_file.read(), "utf-8")
            return text
        
        else:
            st.error(f"Unsupported file type: {file_type}. Please upload PDF, DOCX, or TXT.")
            return None
            
    except Exception as e:
        st.error(f"Error extracting text from resume: {e}")
        return None

def extract_relevant_resume_sections(resume_text):
    """Extract only Experience and Education sections from resume text to reduce token usage in Pass 2 verification"""
    if not resume_text:
        return ""
    
    # Common section headers to look for (case-insensitive)
    experience_keywords = [
        r'experience', r'work experience', r'employment', r'employment history',
        r'professional experience', r'work history', r'career history', r'positions held'
    ]
    education_keywords = [
        r'education', r'academic background', r'academic qualifications',
        r'educational background', r'qualifications', r'degrees'
    ]
    
    # Split resume into lines for easier parsing
    lines = resume_text.split('\n')
    relevant_sections = []
    current_section = None
    in_experience = False
    in_education = False
    
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        if not line_stripped:
            continue
        
        # Check if this line is a section header
        line_lower = line_stripped.lower()
        
        # Check for experience section
        if any(re.search(rf'\b{kw}\b', line_lower) for kw in experience_keywords):
            if not in_experience:
                in_experience = True
                in_education = False
                if current_section:
                    relevant_sections.append(current_section)
                current_section = line + '\n'
            continue
        
        # Check for education section
        if any(re.search(rf'\b{kw}\b', line_lower) for kw in education_keywords):
            if not in_education:
                in_education = True
                if current_section:
                    relevant_sections.append(current_section)
                current_section = line + '\n'
            continue
        
        # Check if we hit another major section (stop collecting)
        major_sections = [r'summary', r'objective', r'skills', r'certifications', 
                         r'awards', r'publications', r'projects', r'contact', r'personal']
        if any(re.search(rf'\b{section}\b', line_lower) for section in major_sections):
            if in_experience or in_education:
                if current_section:
                    relevant_sections.append(current_section)
                current_section = None
                in_experience = False
                in_education = False
            continue
        
        # If we're in a relevant section, add the line
        if in_experience or in_education:
            if current_section:
                current_section += line + '\n'
    
    # Add the last section if we're still collecting
    if current_section and (in_experience or in_education):
        relevant_sections.append(current_section)
    
    # Combine all relevant sections
    result = '\n'.join(relevant_sections)
    
    # If we couldn't extract specific sections, try a simpler approach:
    # Look for date patterns and company names (common in experience sections)
    if not result or len(result) < 100:
        # Fallback: extract text that likely contains experience/education
        # Look for patterns like dates, job titles, company names
        date_pattern = r'\b(19|20)\d{2}\b|\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}'
        # Get context around lines with dates (likely experience/education)
        result_lines = []
        for line in lines:
            if re.search(date_pattern, line, re.IGNORECASE):
                result_lines.append(line)
            elif result_lines:  # Add a few lines after dates for context
                if len([l for l in result_lines[-3:] if l.strip()]) < 3:
                    result_lines.append(line)
                else:
                    break
        if result_lines:
            result = '\n'.join(result_lines[:50])  # Limit to 50 lines
    
    # Limit the result to a reasonable size (much smaller than 4000 chars)
    if result:
        # Limit to approximately 2000 characters (about half of original 4000)
        # This still provides enough context for verification while reducing tokens
        return result[:2000] if len(result) > 2000 else result
    
    # Final fallback: return empty string (will use Pass 1 data only for verification)
    return ""

def extract_profile_from_resume(resume_text):
    """Use Azure OpenAI to extract structured profile information from resume text with two-pass self-correction"""
    try:
        text_gen = get_text_generator()
        
        # Check if text generator was initialized successfully
        if text_gen is None:
            st.error("⚠️ Azure OpenAI is not configured. Please configure AZURE_OPENAI_API_KEY and AZURE_OPENAI_ENDPOINT in your Streamlit secrets.")
            return None
        
        # FIRST PASS: Initial extraction
        prompt_pass1 = f"""You are an expert at parsing resumes. Extract structured information from the following resume text.

RESUME TEXT:
{resume_text}

Please extract and return the following information in JSON format:
{{
    "name": "Full name",
    "email": "Email address",
    "phone": "Phone number",
    "location": "City, State/Country",
    "linkedin": "LinkedIn URL if mentioned",
    "portfolio": "Portfolio/website URL if mentioned",
    "summary": "Professional summary or objective (2-3 sentences)",
    "experience": "Work experience in chronological order with job titles, companies, dates, and key achievements (formatted as bullet points)",
    "education": "Education details including degrees, institutions, and graduation dates",
    "skills": "Comma-separated list of technical and soft skills",
    "certifications": "Professional certifications, awards, publications, or other achievements"
}}

Important:
- If information is not found, use "N/A" or empty string
- Format experience with clear job titles, companies, dates, and bullet points for achievements
- Extract all relevant skills mentioned
- Keep the summary concise but informative
- Return ONLY valid JSON, no additional text or markdown"""
        
        payload_pass1 = {
            "messages": [
                {"role": "system", "content": "You are a resume parser. Extract structured information and return only valid JSON."},
                {"role": "user", "content": prompt_pass1}
            ],
            "max_tokens": 2000,
            "temperature": 0.3,
            "response_format": {"type": "json_object"}
        }
        
        # Send keepalive before LLM API call (can take 10-30s)
        _websocket_keepalive("Extracting profile information...")
        
        def make_request_pass1():
            return requests.post(
                text_gen.url,
                headers=text_gen.headers,
                json=payload_pass1,
                timeout=45
            )
        
        response_pass1 = api_call_with_retry(make_request_pass1, max_retries=3)
        
        if not response_pass1 or response_pass1.status_code != 200:
            # api_call_with_retry already handled 429 retries
            if response_pass1 and response_pass1.status_code == 429:
                st.error("🚫 Rate limit reached for profile extraction after retries. Please wait a few minutes and try again.")
            else:
                error_detail = response_pass1.text[:200] if response_pass1 and response_pass1.text else "No error details"
                endpoint_info = f"Endpoint: {text_gen.url.split('/deployments')[0]}" if text_gen else "Endpoint: Not configured"
                st.error(f"API Error: {response_pass1.status_code if response_pass1 else 'Unknown'} - {error_detail}\n\n{endpoint_info}")
            return None
        
        result_pass1 = response_pass1.json()
        content_pass1 = result_pass1['choices'][0]['message']['content']
        
        # Track token usage for first pass
        if text_gen.token_tracker and 'usage' in result_pass1:
            usage = result_pass1['usage']
            prompt_tokens = usage.get('prompt_tokens', 0)
            completion_tokens = usage.get('completion_tokens', 0)
            text_gen.token_tracker.add_completion_tokens(prompt_tokens, completion_tokens)
        
        # Parse first pass JSON
        try:
            profile_data_pass1 = json.loads(content_pass1)
        except json.JSONDecodeError:
            # Try to find JSON in the response
            json_match = re.search(r'\{.*\}', content_pass1, re.DOTALL)
            if json_match:
                profile_data_pass1 = json.loads(json_match.group())
            else:
                st.error("Could not parse extracted profile data from first pass. Please try again.")
                return None
        
        # SECOND PASS: Self-correction - verify dates and company names
        # This pass is OPTIONAL and disabled by default for faster processing
        # Set ENABLE_PROFILE_PASS2=true in environment to enable (adds 10-30s)
        if not ENABLE_PROFILE_PASS2:
            # Skip Pass 2 for faster processing - return Pass 1 result directly
            return profile_data_pass1
        
        # Extract only relevant sections (Experience and Education) to reduce token usage
        relevant_resume_sections = extract_relevant_resume_sections(resume_text)
        
        # Build verification prompt with only relevant sections
        if relevant_resume_sections:
            resume_context = f"""RELEVANT RESUME SECTIONS (Experience and Education only):
{relevant_resume_sections}"""
        else:
            # Fallback: if we can't extract sections, use a smaller chunk or just Pass 1 data
            # This should rarely happen, but provides a safety net
            resume_context = f"""RELEVANT RESUME SECTIONS (limited):
{resume_text[:1500]}"""
        
        prompt_pass2 = f"""You are a resume quality checker. Review the extracted profile data against the relevant resume sections and verify accuracy, especially for dates and company names.

{resume_context}

EXTRACTED PROFILE DATA (from first pass):
{json.dumps(profile_data_pass1, indent=2)}

Please review and correct the extracted data, paying special attention to:
1. **Dates** - Verify all employment dates, education dates, and certification dates are accurate
2. **Company Names** - Verify all company/organization names are spelled correctly
3. **Job Titles** - Verify job titles are accurate
4. **Education Institutions** - Verify institution names are correct

Return the corrected profile data in the same JSON format. If everything is correct, return the data as-is. If corrections are needed, return the corrected version.

Return ONLY valid JSON with this structure:
{{
    "name": "Full name",
    "email": "Email address",
    "phone": "Phone number",
    "location": "City, State/Country",
    "linkedin": "LinkedIn URL if mentioned",
    "portfolio": "Portfolio/website URL if mentioned",
    "summary": "Professional summary or objective (2-3 sentences)",
    "experience": "Work experience in chronological order with job titles, companies, dates, and key achievements (formatted as bullet points)",
    "education": "Education details including degrees, institutions, and graduation dates",
    "skills": "Comma-separated list of technical and soft skills",
    "certifications": "Professional certifications, awards, publications, or other achievements"
}}

Return ONLY valid JSON, no additional text or markdown."""
        
        payload_pass2 = {
            "messages": [
                {"role": "system", "content": "You are a resume quality checker. Verify and correct extracted data, especially dates and company names. Return only valid JSON."},
                {"role": "user", "content": prompt_pass2}
            ],
            "max_tokens": 2000,
            "temperature": 0.1,  # Lower temperature for more accurate corrections
            "response_format": {"type": "json_object"}
        }
        
        # Send keepalive before second LLM API call
        _websocket_keepalive("Verifying profile data...")
        
        def make_request_pass2():
            return requests.post(
                text_gen.url,
                headers=text_gen.headers,
                json=payload_pass2,
                timeout=45
            )
        
        response_pass2 = api_call_with_retry(make_request_pass2, max_retries=3)
        
        if response_pass2 and response_pass2.status_code == 200:
            result_pass2 = response_pass2.json()
            content_pass2 = result_pass2['choices'][0]['message']['content']
            
            # Track token usage for second pass
            if text_gen.token_tracker and 'usage' in result_pass2:
                usage = result_pass2['usage']
                prompt_tokens = usage.get('prompt_tokens', 0)
                completion_tokens = usage.get('completion_tokens', 0)
                text_gen.token_tracker.add_completion_tokens(prompt_tokens, completion_tokens)
            
            # Parse second pass JSON (corrected version)
            try:
                profile_data_corrected = json.loads(content_pass2)
                return profile_data_corrected
            except json.JSONDecodeError:
                # If second pass fails, return first pass result
                st.warning("⚠️ Self-correction pass failed, using initial extraction. Some details may need manual verification.")
                return profile_data_pass1
        else:
            # If second pass fails, return first pass result
            st.warning("⚠️ Self-correction pass failed, using initial extraction. Some details may need manual verification.")
            return profile_data_pass1
            
    except Exception as e:
        st.error(f"Error extracting profile: {e}")
        return None

def display_user_profile():
    """Display and edit user profile"""
    st.header("👤 Your Profile")
    st.caption("Fill in your information to generate tailored resumes")
    
    # Resume upload section
    st.markdown("---")
    st.subheader("📄 Upload Your Resume (Optional)")
    st.caption("Upload your resume to automatically extract your information")
    
    # Determine available file types based on available libraries
    available_types = ['txt']  # TXT is always available
    if PYPDF2_AVAILABLE:
        available_types.insert(0, 'pdf')
    if DOCX_AVAILABLE:
        available_types.insert(1 if 'pdf' in available_types else 0, 'docx')
    
    uploaded_file = st.file_uploader(
        "Choose a resume file",
        type=available_types,
        help=f"Supported formats: {', '.join(t.upper() for t in available_types)}"
    )
    
    if uploaded_file is not None:
        if st.button("🔍 Extract Information from Resume", type="primary", use_container_width=True):
            # Use progress bar for better UX
            progress_bar = st.progress(0, text="📖 Reading resume...")
            
            # Extract text from resume
            resume_text = extract_text_from_resume(uploaded_file)
            
            if resume_text:
                progress_bar.progress(25, text=f"✅ Read {len(resume_text)} characters")
                # Store resume text for job matching
                st.session_state.resume_text = resume_text
                
                # Show extracted text preview
                with st.expander("📝 Preview Extracted Text"):
                    st.text(resume_text[:1000] + "..." if len(resume_text) > 1000 else resume_text)
                
                # Extract structured information
                progress_bar.progress(35, text="🤖 Extracting profile with AI...")
                profile_data = extract_profile_from_resume(resume_text)
                
                if profile_data:
                    progress_bar.progress(75, text="📊 Finalizing profile...")
                    # Update session state with extracted data
                    st.session_state.user_profile = {
                        'name': profile_data.get('name', ''),
                        'email': profile_data.get('email', ''),
                        'phone': profile_data.get('phone', ''),
                        'location': profile_data.get('location', ''),
                        'linkedin': profile_data.get('linkedin', ''),
                        'portfolio': profile_data.get('portfolio', ''),
                        'summary': profile_data.get('summary', ''),
                        'experience': profile_data.get('experience', ''),
                        'education': profile_data.get('education', ''),
                        'skills': profile_data.get('skills', ''),
                        'certifications': profile_data.get('certifications', '')
                    }
                    
                    # Generate resume embedding AFTER profile extraction
                    progress_bar.progress(90, text="🔗 Creating search embedding...")
                    generate_and_store_resume_embedding(resume_text, st.session_state.user_profile)
                    
                    progress_bar.progress(100, text="✅ Complete!")
                    time.sleep(0.3)
                    progress_bar.empty()
                    
                    st.success("✅ Profile information extracted successfully! Review and edit below.")
                    st.balloons()
                    time.sleep(0.5)
                    st.rerun()
                else:
                    progress_bar.empty()
                    st.warning("⚠️ Could not extract structured information. Please fill in manually.")
            else:
                progress_bar.empty()
                st.error("❌ Could not read the resume file. Please try a different file.")
    
    st.markdown("---")
    
    with st.form("user_profile_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            name = st.text_input("Full Name", value=st.session_state.user_profile.get('name', ''))
            email = st.text_input("Email", value=st.session_state.user_profile.get('email', ''))
            phone = st.text_input("Phone", value=st.session_state.user_profile.get('phone', ''))
        
        with col2:
            location = st.text_input("Location", value=st.session_state.user_profile.get('location', ''))
            linkedin = st.text_input("LinkedIn URL", value=st.session_state.user_profile.get('linkedin', ''))
            portfolio = st.text_input("Portfolio/Website", value=st.session_state.user_profile.get('portfolio', ''))
        
        summary = st.text_area(
            "Professional Summary",
            value=st.session_state.user_profile.get('summary', ''),
            height=100,
            placeholder="Brief overview of your professional background..."
        )
        
        experience = st.text_area(
            "Work Experience",
            value=st.session_state.user_profile.get('experience', ''),
            height=150,
            placeholder="List your work experience with job titles, companies, dates, and key achievements..."
        )
        
        education = st.text_area(
            "Education",
            value=st.session_state.user_profile.get('education', ''),
            height=100,
            placeholder="Degrees, institutions, graduation dates..."
        )
        
        skills = st.text_area(
            "Skills",
            value=st.session_state.user_profile.get('skills', ''),
            height=80,
            placeholder="List your technical and soft skills (comma-separated)..."
        )
        
        certifications = st.text_area(
            "Certifications & Awards",
            value=st.session_state.user_profile.get('certifications', ''),
            height=80,
            placeholder="Professional certifications, awards, publications..."
        )
        
        submitted = st.form_submit_button("💾 Save Profile", use_container_width=True, type="primary")
        
        if submitted:
            st.session_state.user_profile = {
                'name': name,
                'email': email,
                'phone': phone,
                'location': location,
                'linkedin': linkedin,
                'portfolio': portfolio,
                'summary': summary,
                'experience': experience,
                'education': education,
                'skills': skills,
                'certifications': certifications
            }
            st.success("✅ Profile saved successfully!")
            time.sleep(1)
            st.rerun()

def render_structured_resume_editor(resume_data):
    """Render structured resume JSON in editable Streamlit form"""
    if not resume_data:
        return None
    
    edited_data = {}
    
    st.subheader("📋 Your Tailored Resume")
    st.caption("Edit the sections below to customize your resume")
    
    # Header Section
    with st.expander("👤 Header Information", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            edited_data['header'] = {
                'name': st.text_input("Full Name", value=resume_data.get('header', {}).get('name', ''), key='resume_name'),
                'title': st.text_input("Professional Title", value=resume_data.get('header', {}).get('title', ''), key='resume_title'),
                'email': st.text_input("Email", value=resume_data.get('header', {}).get('email', ''), key='resume_email'),
                'phone': st.text_input("Phone", value=resume_data.get('header', {}).get('phone', ''), key='resume_phone'),
            }
        with col2:
            edited_data['header']['location'] = st.text_input("Location", value=resume_data.get('header', {}).get('location', ''), key='resume_location')
            edited_data['header']['linkedin'] = st.text_input("LinkedIn URL", value=resume_data.get('header', {}).get('linkedin', ''), key='resume_linkedin')
            edited_data['header']['portfolio'] = st.text_input("Portfolio URL", value=resume_data.get('header', {}).get('portfolio', ''), key='resume_portfolio')
    
    # Summary
    col_summary1, col_summary2 = st.columns([4, 1])
    with col_summary1:
        edited_data['summary'] = st.text_area(
            "Professional Summary",
            value=resume_data.get('summary', ''),
            height=100,
            key='resume_summary'
        )
    with col_summary2:
        if st.button("✨ Refine with AI", key='refine_summary', use_container_width=True, help="Use AI to improve this section"):
            with st.spinner("🤖 Refining summary..."):
                text_gen = get_text_generator()
                if text_gen is None:
                    st.error("⚠️ Azure OpenAI is not configured. Please configure AZURE_OPENAI_API_KEY and AZURE_OPENAI_ENDPOINT in your Streamlit secrets.")
                    return
                refinement_prompt = f"""Improve this professional summary. Make it more impactful, quantified, and tailored. Keep it concise (2-3 sentences).

Current Summary:
{edited_data.get('summary', resume_data.get('summary', ''))}

Return ONLY the improved summary text, no additional explanation."""
                
                payload = {
                    "messages": [
                        {"role": "system", "content": "You are a resume writing expert. Improve professional summaries to be more impactful and quantified."},
                        {"role": "user", "content": refinement_prompt}
                    ],
                    "max_tokens": 200,
                    "temperature": 0.7
                }
                
                def make_request():
                    return requests.post(text_gen.url, headers=text_gen.headers, json=payload, timeout=30)
                
                response = api_call_with_retry(make_request, max_retries=2)
                if response and response.status_code == 200:
                    result = response.json()
                    refined_text = result['choices'][0]['message']['content'].strip()
                    # Update the text area value
                    st.session_state['resume_summary'] = refined_text
                    st.rerun()
    
    # Skills
    skills_list = resume_data.get('skills_highlighted', [])
    skills_text = ', '.join(skills_list) if skills_list else ''
    skills_input = st.text_area(
        "Highlighted Skills (comma-separated)",
        value=skills_text,
        height=60,
        key='resume_skills',
        help="List skills separated by commas"
    )
    edited_data['skills_highlighted'] = [s.strip() for s in skills_input.split(',') if s.strip()]
    
    # Experience
    st.subheader("💼 Work Experience")
    edited_data['experience'] = []
    
    experience_list = resume_data.get('experience', [])
    for i, exp in enumerate(experience_list):
        with st.expander(f"📌 {exp.get('company', 'Company')} - {exp.get('title', 'Position')}", expanded=(i == 0)):
            col1, col2 = st.columns([2, 1])
            with col1:
                company = st.text_input("Company", value=exp.get('company', ''), key=f'exp_company_{i}')
                title = st.text_input("Job Title", value=exp.get('title', ''), key=f'exp_title_{i}')
            with col2:
                dates = st.text_input("Date Range", value=exp.get('dates', ''), key=f'exp_dates_{i}')
            
            st.write("**Key Achievements:**")
            bullets = exp.get('bullets', [])
            edited_bullets = []
            for j, bullet in enumerate(bullets):
                col_bullet1, col_bullet2 = st.columns([4, 1])
                with col_bullet1:
                    bullet_text = st.text_area(
                        f"Bullet {j+1}",
                        value=bullet,
                        height=60,
                        key=f'exp_bullet_{i}_{j}'
                    )
                with col_bullet2:
                    if st.button("✨", key=f'refine_bullet_{i}_{j}', help="Refine this bullet with AI", use_container_width=True):
                        with st.spinner("🤖 Refining..."):
                            text_gen = get_text_generator()
                            if text_gen is None:
                                st.error("⚠️ Azure OpenAI is not configured. Please configure AZURE_OPENAI_API_KEY and AZURE_OPENAI_ENDPOINT in your Streamlit secrets.")
                                return
                            refinement_prompt = f"""Improve this resume bullet point. Make it more quantified, impactful, and achievement-focused. Use numbers, percentages, or metrics when possible.

Current Bullet:
{bullet_text if bullet_text else bullet}

Return ONLY the improved bullet point, no additional text."""
                            
                            payload = {
                                "messages": [
                                    {"role": "system", "content": "You are a resume writing expert. Improve bullet points to be quantified and achievement-focused."},
                                    {"role": "user", "content": refinement_prompt}
                                ],
                                "max_tokens": 150,
                                "temperature": 0.7
                            }
                            
                            def make_request():
                                return requests.post(text_gen.url, headers=text_gen.headers, json=payload, timeout=30)
                            
                            response = api_call_with_retry(make_request, max_retries=2)
                            if response and response.status_code == 200:
                                result = response.json()
                                refined_text = result['choices'][0]['message']['content'].strip()
                                st.session_state[f'exp_bullet_{i}_{j}'] = refined_text
                                st.rerun()
                
                if bullet_text.strip():
                    edited_bullets.append(bullet_text.strip())
            
            # Allow adding new bullets
            if st.button(f"➕ Add Bullet Point", key=f'add_bullet_{i}'):
                edited_bullets.append("")
                st.rerun()
            
            edited_data['experience'].append({
                'company': company,
                'title': title,
                'dates': dates,
                'bullets': edited_bullets
            })
    
    # Education
    edited_data['education'] = st.text_area(
        "Education",
        value=resume_data.get('education', ''),
        height=100,
        key='resume_education'
    )
    
    # Certifications
    edited_data['certifications'] = st.text_area(
        "Certifications & Awards",
        value=resume_data.get('certifications', ''),
        height=100,
        key='resume_certifications'
    )
    
    return edited_data

def generate_docx_from_json(resume_data, filename="resume.docx"):
    """Generate a professional .docx file from structured resume JSON"""
    # Check if python-docx is available
    if not DOCX_AVAILABLE:
        st.error("DOCX generation is not available. Please download as TXT instead.")
        return None
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Header Section
        header = resume_data.get('header', {})
        if header.get('name'):
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(header['name'])
            name_run.font.size = Pt(18)
            name_run.font.bold = True
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact information
        contact_info = []
        if header.get('email'):
            contact_info.append(header['email'])
        if header.get('phone'):
            contact_info.append(header['phone'])
        if header.get('location'):
            contact_info.append(header['location'])
        if header.get('linkedin'):
            contact_info.append(header['linkedin'])
        if header.get('portfolio'):
            contact_info.append(header['portfolio'])
        
        if contact_info:
            contact_para = doc.add_paragraph(' | '.join(contact_info))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.runs[0].font.size = Pt(10)
        
        doc.add_paragraph()  # Spacing
        
        # Professional Title
        if header.get('title'):
            title_para = doc.add_paragraph(header['title'])
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.runs[0].font.size = Pt(12)
            title_para.runs[0].italic = True
            doc.add_paragraph()  # Spacing
        
        # Summary
        if resume_data.get('summary'):
            doc.add_heading('Professional Summary', level=2)
            summary_para = doc.add_paragraph(resume_data['summary'])
            summary_para.runs[0].font.size = Pt(11)
            doc.add_paragraph()  # Spacing
        
        # Skills
        skills = resume_data.get('skills_highlighted', [])
        if skills:
            doc.add_heading('Key Skills', level=2)
            skills_text = ' • '.join(skills)
            skills_para = doc.add_paragraph(skills_text)
            skills_para.runs[0].font.size = Pt(11)
            doc.add_paragraph()  # Spacing
        
        # Experience
        experience = resume_data.get('experience', [])
        if experience:
            doc.add_heading('Professional Experience', level=2)
            for exp in experience:
                # Company and Title
                exp_header = doc.add_paragraph()
                exp_header.add_run(exp.get('title', '')).bold = True
                if exp.get('company'):
                    exp_header.add_run(f" at {exp['company']}")
                if exp.get('dates'):
                    exp_header.add_run(f" | {exp['dates']}").italic = True
                
                # Bullet points
                bullets = exp.get('bullets', [])
                for bullet in bullets:
                    if bullet.strip():
                        bullet_para = doc.add_paragraph(bullet, style='List Bullet')
                        bullet_para.runs[0].font.size = Pt(10)
                
                doc.add_paragraph()  # Spacing between experiences
        
        # Education
        if resume_data.get('education'):
            doc.add_heading('Education', level=2)
            edu_para = doc.add_paragraph(resume_data['education'])
            edu_para.runs[0].font.size = Pt(11)
            doc.add_paragraph()  # Spacing
        
        # Certifications
        if resume_data.get('certifications'):
            doc.add_heading('Certifications & Awards', level=2)
            cert_para = doc.add_paragraph(resume_data['certifications'])
            cert_para.runs[0].font.size = Pt(11)
        
        # Save to BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io
        
    except Exception as e:
        st.error(f"Error generating DOCX: {e}")
        return None

def display_match_score_feedback(match_score, missing_keywords, job_title):
    """Display match score and feedback to user"""
    if match_score is None:
        return
    
    st.markdown("---")
    st.subheader("🎯 Resume Match Analysis")
    
    # Match score with color coding
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        score_percent = match_score * 100
        
        if score_percent >= 80:
            score_color = "🟢"
            feedback = "Excellent match! Your resume aligns well with this position."
        elif score_percent >= 60:
            score_color = "🟡"
            feedback = "Good match. Consider adding more relevant keywords."
        else:
            score_color = "🔴"
            feedback = "Moderate match. Your resume may need more tailoring."
        
        st.markdown(f"""
        <div style="text-align: center; margin: 1rem 0;">
            <div class="match-score-display">{score_percent:.0f}%</div>
            <p style="color: var(--text-secondary); margin-top: 0.5rem;">Match Score</p>
        </div>
        """, unsafe_allow_html=True)
        st.caption(f"**Analysis:** {feedback}")
    
    # Missing keywords
    if missing_keywords:
        st.warning(f"**Missing Keywords:** {', '.join(missing_keywords[:5])}")
        if len(missing_keywords) > 5:
            with st.expander(f"See all {len(missing_keywords)} missing keywords"):
                st.write(', '.join(missing_keywords))
        
        st.info("💡 **Tip:** Consider adding these keywords to your resume if you have experience with them. Be honest - only include skills you actually possess.")
    else:
        st.success("✅ Great! Your resume includes the key keywords from the job description.")
    
    # Feedback
    st.caption(f"**Analysis:** {feedback}")

def display_resume_generator():
    """Display the resume generator interface with structured resume editing"""
    if st.session_state.selected_job is None:
        st.warning("No job selected. Please select a job first.")
        if st.button("← Back to Jobs"):
            st.session_state.show_resume_generator = False
            st.rerun()
        return
    
    job = st.session_state.selected_job
    
    st.markdown('<h1 class="main-header">📄 Resume Generator</h1>', unsafe_allow_html=True)
    
    # Display selected job info
    st.markdown(f"""
    <div class="job-card">
        <h3 style="color: var(--primary-accent); margin: 0;">{job['title']}</h3>
        <p style="margin: 0.5rem 0; color: var(--text-secondary);">🏢 {job['company']} • 📍 {job['location']}</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Check if user profile is complete
    if not st.session_state.user_profile.get('name') or not st.session_state.user_profile.get('experience'):
        st.error("⚠️ Please complete your profile first!")
        if st.button("← Go to Profile"):
            st.session_state.show_resume_generator = False
            st.rerun()
        return
    
    col1, col2 = st.columns([3, 1])
    
    with col1:
        st.write("**Your Profile:**", st.session_state.user_profile.get('name', 'N/A'))
    
    with col2:
        if st.button("← Back to Jobs"):
            st.session_state.show_resume_generator = False
            st.session_state.generated_resume = None
            st.session_state.match_score = None
            st.session_state.missing_keywords = None
            st.rerun()
    
    st.markdown("---")
    
    # Generate resume button
    if st.button("🚀 Generate Tailored Resume", type="primary", use_container_width=True):
        with st.spinner("🤖 Creating your personalized resume using AI..."):
            text_gen = get_text_generator()
            if text_gen is None:
                st.error("⚠️ Azure OpenAI is not configured. Please configure AZURE_OPENAI_API_KEY and AZURE_OPENAI_ENDPOINT in your Streamlit secrets.")
                return
            # Get raw resume text if available
            raw_resume_text = st.session_state.get('resume_text')
            resume_data = text_gen.generate_resume(
                st.session_state.user_profile, 
                job,
                raw_resume_text=raw_resume_text
            )
            
            if resume_data:
                st.session_state.generated_resume = resume_data
                
                # Calculate match score
                with st.spinner("📊 Analyzing resume match..."):
                    embedding_gen = get_embedding_generator()
                    # Convert resume JSON to text for comparison
                    resume_text = json.dumps(resume_data, indent=2)
                    match_score, missing_keywords = text_gen.calculate_match_score(
                        resume_text,
                        job.get('description', ''),
                        embedding_gen
                    )
                    st.session_state.match_score = match_score
                    st.session_state.missing_keywords = missing_keywords
                
                st.success("✅ Resume generated successfully!")
                st.balloons()
                time.sleep(0.5)
                st.rerun()
            else:
                st.error("❌ Failed to generate resume. Please try again.")
    
    # Display match score if available
    if st.session_state.generated_resume and st.session_state.get('match_score') is not None:
        display_match_score_feedback(
            st.session_state.match_score,
            st.session_state.missing_keywords,
            job['title']
        )
    
    # Display generated resume in structured form
    if st.session_state.generated_resume:
        st.markdown("---")
        
        # Render structured editor
        edited_resume_data = render_structured_resume_editor(st.session_state.generated_resume)
        
        # Update session state with edited data
        if edited_resume_data:
            st.session_state.generated_resume = edited_resume_data
        
        st.markdown("---")
        
        # Download buttons
        col1, col2, col3, col4, col5 = st.columns(5)
        
        with col1:
            # Download as PDF
            pdf_file = generate_pdf_from_json(
                st.session_state.generated_resume,
                filename=f"resume_{job['company']}_{job['title']}.pdf"
            )
            if pdf_file:
                st.download_button(
                    label="📥 Download as PDF",
                    data=pdf_file,
                    file_name=f"resume_{job['company']}_{job['title']}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
        
        with col2:
            # Download as DOCX
            docx_file = generate_docx_from_json(
                st.session_state.generated_resume,
                filename=f"resume_{job['company']}_{job['title']}.docx"
            )
            if docx_file:
                st.download_button(
                    label="📥 Download as DOCX",
                    data=docx_file,
                    file_name=f"resume_{job['company']}_{job['title']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        
        with col3:
            # Download as JSON
            json_data = json.dumps(st.session_state.generated_resume, indent=2)
            st.download_button(
                label="📥 Download as JSON",
                data=json_data,
                file_name=f"resume_{job['company']}_{job['title']}.json",
                mime="application/json",
                use_container_width=True
            )
        
        with col4:
            # Download as TXT (formatted text version)
            txt_content = format_resume_as_text(st.session_state.generated_resume)
            st.download_button(
                label="📥 Download as TXT",
                data=txt_content,
                file_name=f"resume_{job['company']}_{job['title']}.txt",
                mime="text/plain",
                use_container_width=True
            )
        
        with col5:
            # Apply to job button
            if job['url'] != '#':
                st.link_button(
                    "🚀 Apply to Job",
                    job['url'],
                    use_container_width=True,
                    type="primary"
                )
        
        # Recalculate match score button
        if st.button("🔄 Recalculate Match Score", use_container_width=True):
            with st.spinner("📊 Recalculating match score..."):
                text_gen = get_text_generator()
                if text_gen is None:
                    st.error("⚠️ Azure OpenAI is not configured. Please configure AZURE_OPENAI_API_KEY and AZURE_OPENAI_ENDPOINT in your Streamlit secrets.")
                    return
                embedding_gen = get_embedding_generator()
                resume_text = json.dumps(st.session_state.generated_resume, indent=2)
                match_score, missing_keywords = text_gen.calculate_match_score(
                    resume_text,
                    job.get('description', ''),
                    embedding_gen
                )
                st.session_state.match_score = match_score
                st.session_state.missing_keywords = missing_keywords
                st.rerun()

def render_sidebar():
    """Render CareerLens sidebar with resume upload, market filters, and analyze button"""
    with st.sidebar:
        # Theme detection is handled globally at app initialization (see app.py line ~977)
        # No need to duplicate the theme script here
        
        # Header with icon and title (Navy theme - white text on dark sidebar)
        st.markdown("""
        <div style="margin-bottom: 2rem;">
            <h2 style="color: white !important; margin-bottom: 0.5rem; display: flex; align-items: center; gap: 0.5rem;">
                🔍 CareerLens
            </h2>
            <p style="color: #94a3b8 !important; font-size: 0.9rem; margin: 0;">AI Career Copilot for Hong Kong</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Resume Upload Section
        st.markdown("---")
        st.markdown("### 1. Upload your CV to begin")
        
        # Determine available file types based on available libraries
        sidebar_available_types = ['txt']  # TXT always available
        if PYPDF2_AVAILABLE:
            sidebar_available_types.insert(0, 'pdf')
        if DOCX_AVAILABLE:
            sidebar_available_types.insert(1 if 'pdf' in sidebar_available_types else 0, 'docx')
        
        uploaded_file = st.file_uploader(
            "Upload your resume",
            type=sidebar_available_types,
            help=f"Supported: {', '.join(t.upper() for t in sidebar_available_types)}. We parse your skills and experience to benchmark you against the market.",
            key="careerlens_resume_upload",
            label_visibility="collapsed"
        )
        
        if uploaded_file is not None:
            # Check if this is a new file upload (different from what's cached)
            file_key = f"{uploaded_file.name}_{uploaded_file.size}"
            current_cached_key = st.session_state.get('_last_uploaded_file_key')
            
            if current_cached_key != file_key:
                # New file uploaded - process it
                progress_bar = st.progress(0, text="📖 Reading resume...")
                resume_text = extract_text_from_resume(uploaded_file)
                
                if resume_text:
                    progress_bar.progress(30, text="✅ Resume read successfully")
                    st.session_state.resume_text = resume_text
                    st.session_state._last_uploaded_file_key = file_key
                    
                    # Extract structured information (Pass 1 only by default - faster!)
                    progress_bar.progress(40, text="🤖 Extracting profile with AI...")
                    profile_data = extract_profile_from_resume(resume_text)
                    
                    if profile_data:
                        progress_bar.progress(80, text="📊 Finalizing profile...")
                        st.session_state.user_profile = {
                            'name': profile_data.get('name', ''),
                            'email': profile_data.get('email', ''),
                            'phone': profile_data.get('phone', ''),
                            'location': profile_data.get('location', ''),
                            'linkedin': profile_data.get('linkedin', ''),
                            'portfolio': profile_data.get('portfolio', ''),
                            'summary': profile_data.get('summary', ''),
                            'experience': profile_data.get('experience', ''),
                            'education': profile_data.get('education', ''),
                            'skills': profile_data.get('skills', ''),
                            'certifications': profile_data.get('certifications', '')
                        }
                        
                        # Generate resume embedding AFTER profile extraction (uses profile for better embedding)
                        progress_bar.progress(90, text="🔗 Creating search embedding...")
                        generate_and_store_resume_embedding(resume_text, st.session_state.user_profile)
                        
                        progress_bar.progress(100, text="✅ Profile ready!")
                        time.sleep(0.3)  # Brief pause to show completion
                        progress_bar.empty()
                        st.success("✅ Profile extracted!")
                    else:
                        progress_bar.empty()
                        st.warning("⚠️ Could not extract profile. Please try again.")
                else:
                    progress_bar.empty()
                    st.error("❌ Could not read the resume file.")
            else:
                # Same file already processed - show cached status
                if st.session_state.user_profile.get('name'):
                    st.success(f"✅ Using profile for: {st.session_state.user_profile.get('name', 'Unknown')}")
        
        # Search Criteria Section
        st.markdown("---")
        st.markdown("### 2. Set Search Criteria")
        
        # Target Domains
        target_domains = st.multiselect(
            "Target Domains",
            options=["FinTech", "ESG & Sustainability", "Data Analytics", "Digital Transformation", 
                    "Investment Banking", "Consulting", "Technology", "Healthcare", "Education"],
            default=st.session_state.get('target_domains', []),
            help="Select industries/domains to search for jobs",
            key="sidebar_target_domains"
        )
        st.session_state.target_domains = target_domains
        
        # Salary Expectation
        salary_expectation = st.slider(
            "Min. Monthly Salary (HKD)",
            min_value=0,
            max_value=150000,
            value=st.session_state.get('salary_expectation', 0),
            step=5000,
            help="Set to 0 to disable salary filtering",
            key="sidebar_salary"
        )
        st.session_state.salary_expectation = salary_expectation
        
        # Primary Action Button
        st.markdown("---")
        analyze_button = st.button(
            "Analyze Profile & Find Matches",
            type="primary",
            use_container_width=True,
            key="careerlens_analyze"
        )
        
        if analyze_button:
            # Fetch jobs and perform matching
            if not st.session_state.resume_text and not st.session_state.user_profile.get('summary'):
                st.error("⚠️ Please upload your CV first!")
            else:
                # Use user-specified filters from sidebar
                target_domains = st.session_state.get('target_domains', [])
                salary_expectation = st.session_state.get('salary_expectation', 0)
                
                # Build search query from user-specified domains
                search_query = " ".join(target_domains) if target_domains else "Hong Kong jobs"
                scraper = get_job_scraper()
                
                if scraper is None:
                    st.error("⚠️ Job scraper not configured. Please check your RAPIDAPI_KEY in Streamlit secrets.")
                    return
                
                # Use progress bar for better UX
                progress_bar = st.progress(0, text="🔍 Starting job search...")
                
                # Fetch jobs with cache
                progress_bar.progress(10, text="📡 Fetching jobs from Indeed...")
                jobs = fetch_jobs_with_cache(
                    scraper,
                    search_query,
                    location="Hong Kong",
                    max_rows=25,
                    job_type="fulltime",
                    country="hk",
                    force_refresh=False
                )
                
                if not jobs:
                    progress_bar.empty()
                    st.error("❌ No jobs found from Indeed. Please check your API configuration or try different search criteria.")
                    return
                
                # Show how many jobs were fetched before filtering
                total_fetched = len(jobs)
                progress_bar.progress(30, text=f"✅ Found {total_fetched} jobs, applying filters...")
                
                # Apply domain filters (only if user specified domains)
                if target_domains:
                    jobs = filter_jobs_by_domains(jobs, target_domains)
                
                # Apply salary filter (only if user specified a salary > 0)
                if salary_expectation > 0:
                    jobs = filter_jobs_by_salary(jobs, salary_expectation)
                
                if not jobs:
                    progress_bar.empty()
                    st.warning(f"⚠️ No jobs match your filters. Found {total_fetched} jobs but none passed your criteria. Try reducing salary or selecting different domains.")
                    return
                
                progress_bar.progress(40, text=f"📊 Analyzing {len(jobs)} matching jobs...")
                
                # Perform semantic matching
                embedding_gen = get_embedding_generator()
                if embedding_gen is None:
                    progress_bar.empty()
                    st.error("⚠️ Azure OpenAI is not configured. Please configure AZURE_OPENAI_API_KEY and AZURE_OPENAI_ENDPOINT in your Streamlit secrets.")
                    return
                
                desired_matches = min(15, len(jobs))
                jobs_to_index_limit = _determine_index_limit(len(jobs), desired_matches)
                top_match_count = min(desired_matches, jobs_to_index_limit)
                search_engine = SemanticJobSearch(embedding_gen)
                
                progress_bar.progress(50, text=f"🔗 Creating job embeddings ({jobs_to_index_limit} jobs)...")
                search_engine.index_jobs(jobs, max_jobs_to_index=jobs_to_index_limit)
                
                # Use pre-computed resume embedding if available
                resume_embedding = st.session_state.get('resume_embedding')
                if not resume_embedding and st.session_state.resume_text:
                    progress_bar.progress(70, text="🔗 Creating resume embedding...")
                    resume_embedding = generate_and_store_resume_embedding(
                        st.session_state.resume_text,
                        st.session_state.user_profile if st.session_state.user_profile else None
                    )
                
                # Fallback: build query string if no resume embedding available
                resume_query = None
                if not resume_embedding:
                    if st.session_state.resume_text:
                        resume_query = st.session_state.resume_text
                        if st.session_state.user_profile.get('summary'):
                            profile_data = f"{st.session_state.user_profile.get('summary', '')} {st.session_state.user_profile.get('experience', '')} {st.session_state.user_profile.get('skills', '')}"
                            resume_query = f"{resume_query} {profile_data}"
                    else:
                        resume_query = f"{st.session_state.user_profile.get('summary', '')} {st.session_state.user_profile.get('experience', '')} {st.session_state.user_profile.get('skills', '')} {st.session_state.user_profile.get('education', '')}"
                
                progress_bar.progress(80, text="🎯 Finding best matches...")
                results = search_engine.search(query=resume_query, top_k=top_match_count, resume_embedding=resume_embedding)
                
                if results:
                    progress_bar.progress(90, text="📈 Calculating skill matches...")
                    # Calculate skill matches (now uses fast string matching by default)
                    user_skills = st.session_state.user_profile.get('skills', '')
                    for i, result in enumerate(results):
                        job_skills = result['job'].get('skills', [])
                        skill_score, missing_skills = search_engine.calculate_skill_match(user_skills, job_skills)
                        result['skill_match_score'] = skill_score
                        result['missing_skills'] = missing_skills
                        
                        # Calculate combined match score (weighted: 60% semantic, 40% skill)
                        semantic_score = result.get('similarity_score', 0.0)
                        combined_score = (semantic_score * 0.6) + (skill_score * 0.4)
                        result['combined_match_score'] = combined_score
                    
                    # Sort results by combined match score (highest to lowest)
                    results.sort(key=lambda x: x.get('combined_match_score', 0.0), reverse=True)
                    
                    progress_bar.progress(100, text="✅ Analysis complete!")
                    time.sleep(0.3)
                    progress_bar.empty()
                    
                    st.session_state.matched_jobs = results
                    st.session_state.dashboard_ready = True
                    
                    # Clean up memory after analysis
                    gc.collect()
                    
                    st.rerun()
                else:
                    progress_bar.empty()
                    st.error("❌ No matching jobs found. Please try different filters.")
        
        # Skill Matching Calculation Matrix
        display_skill_matching_matrix(st.session_state.user_profile)

def display_skill_matching_matrix(user_profile):
    """Display skill matching calculation matrix to help users understand ranking"""
    st.markdown("---")
    st.markdown("### 📊 How Job Ranking Works")
    
    user_skills = user_profile.get('skills', '') if user_profile else ''
    
    if not user_skills:
        st.info("💡 **Skill-Based Ranking**: Jobs are ranked purely by how many required skills you match. Upload your profile to see your skills analyzed.")
        return
    
    # Parse user skills
    user_skills_list = [s.strip() for s in str(user_skills).split(',') if s.strip()]
    
    if not user_skills_list:
        st.info("💡 **Skill-Based Ranking**: Jobs are ranked purely by how many required skills you match.")
        return
    
    # Show user's skills
    st.markdown("#### Your Skills")
    skills_display = ", ".join(user_skills_list[:10])
    if len(user_skills_list) > 10:
        skills_display += f" (+{len(user_skills_list) - 10} more)"
    st.markdown(f"**{len(user_skills_list)} skills identified:** {skills_display}")
    
    st.markdown("---")
    
    # Explanation of calculation
    st.markdown("#### Ranking Formula")
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown("""
        **Skill Match Score =**
        
        ```
        Matched Skills
        ──────────────
        Required Skills
        ```
        
        **Example:**
        - Job requires: Python, SQL, React, Docker
        - You have: Python, SQL, React
        - **Score: 3/4 = 75%**
        """)
    
    with col2:
        st.markdown("""
        **Ranking Logic:**
        
        1. ✅ Jobs are fetched from job boards
        2. 🔍 Your skills are matched against each job's required skills
        3. 📊 Jobs are sorted by skill match score (highest first)
        4. 🎯 Top matches appear at the top of the list
        
        **Why this approach?**
        - Transparent: You see exactly why jobs rank high
        - Objective: Based on concrete skill requirements
        - Actionable: Shows which skills to learn
        """)
    
    st.markdown("---")
    
    # Show matching method
    st.markdown("#### Matching Method")
    
    method_col1, method_col2 = st.columns(2)
    
    with method_col1:
        st.markdown("""
        **Semantic Matching** (Primary)
        - Uses AI embeddings to understand skill similarity
        - Recognizes related skills (e.g., "JavaScript" ≈ "JS")
        - Handles variations and synonyms
        - Threshold: 70% similarity required
        """)
    
    with method_col2:
        st.markdown("""
        **String Matching** (Fallback)
        - Used when semantic matching unavailable
        - Direct text comparison
        - Case-insensitive matching
        - Handles partial matches
        """)
    
    # Show example calculation if we have matched jobs
    if 'matched_jobs' in st.session_state and st.session_state.matched_jobs:
        st.markdown("---")
        st.markdown("#### Example: Top Match Breakdown")
        
        top_match = st.session_state.matched_jobs[0] if st.session_state.matched_jobs else None
        if top_match:
            job = top_match.get('job', {})
            job_skills = job.get('skills', [])
            skill_score = top_match.get('skill_match_score', 0.0)
            matched_count = int(skill_score * len(job_skills)) if job_skills else 0
            
            if job_skills:
                st.markdown(f"**{job.get('title', 'Job')} at {job.get('company', 'Company')}**")
                st.markdown(f"**Match Score: {int(skill_score * 100)}%** ({matched_count}/{len(job_skills)} skills matched)")
                
                # Show which skills matched
                job_skills_lower = [s.lower().strip() for s in job_skills if isinstance(s, str)]
                user_skills_lower = [s.lower().strip() for s in user_skills_list]
                
                matched_skills_list = []
                missing_skills_list = []
                
                for js in job_skills_lower:
                    matched = False
                    for us in user_skills_lower:
                        if js in us or us in js:
                            matched_skills_list.append(js)
                            matched = True
                            break
                    if not matched:
                        missing_skills_list.append(js)
                
                if matched_skills_list:
                    st.success(f"✅ **Matched Skills:** {', '.join(matched_skills_list[:5])}")
                if missing_skills_list:
                    st.warning(f"⚠️ **Missing Skills:** {', '.join(missing_skills_list[:5])}")

def display_market_positioning_profile(matched_jobs, user_profile):
    """Display Dashboard with 3 key metric cards: Match Score, Est. Salary, Skill Gaps"""
    if not matched_jobs:
        return
    
    # --- METRIC 1: Average Match Score ---
    avg_match_score = sum(r.get('combined_match_score', 0) for r in matched_jobs) / len(matched_jobs)
    match_score_pct = int(avg_match_score * 100)
    
    # Determine match quality delta
    if match_score_pct >= 80:
        match_delta = "Excellent fit"
        match_delta_color = "normal"
    elif match_score_pct >= 60:
        match_delta = "Good fit"
        match_delta_color = "off"
    else:
        match_delta = "Room to improve"
        match_delta_color = "inverse"
    
    # --- METRIC 2: Estimated Salary Band ---
    salary_min, salary_max = calculate_salary_band(matched_jobs)
    avg_salary = (salary_min + salary_max) // 2
    
    # Calculate salary delta (compare with user's expectation if available)
    user_salary_expectation = st.session_state.get('salary_expectation', 0)
    if user_salary_expectation > 0:
        salary_delta_pct = ((avg_salary - user_salary_expectation) / user_salary_expectation * 100) if user_salary_expectation > 0 else 0
        if salary_delta_pct > 0:
            salary_delta = f"+{salary_delta_pct:.0f}% vs target"
        elif salary_delta_pct < 0:
            salary_delta = f"{salary_delta_pct:.0f}% vs target"
        else:
            salary_delta = "Matches target"
    else:
        salary_delta = "Market rate"
    
    # --- METRIC 3: Skill Gaps Count ---
    user_skills = user_profile.get('skills', '')
    all_job_skills = []
    for result in matched_jobs:
        all_job_skills.extend(result['job'].get('skills', []))
    
    # Find missing skills (unique gaps)
    user_skills_list = [s.lower().strip() for s in str(user_skills).split(',') if s.strip()]
    skill_gaps = set()
    for job_skill in all_job_skills:
        if isinstance(job_skill, str):
            job_skill_lower = job_skill.lower().strip()
            if job_skill_lower and not any(us in job_skill_lower or job_skill_lower in us for us in user_skills_list):
                skill_gaps.add(job_skill_lower)
    
    num_skill_gaps = len(skill_gaps)
    
    # Determine skill gap status
    if num_skill_gaps <= 3:
        gap_delta = "Well positioned"
        gap_delta_color = "normal"
    elif num_skill_gaps <= 7:
        gap_delta = "Some upskilling needed"
        gap_delta_color = "off"
    else:
        gap_delta = "Focus on learning"
        gap_delta_color = "inverse"
    
    # --- RENDER 3 DASHBOARD METRIC CARDS ---
    st.markdown("### 📊 Your Dashboard")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        <div class="dashboard-metric-card">
            <div class="dashboard-metric-label">Match Score</div>
            <div class="dashboard-metric-value">{0}%</div>
        </div>
        """.format(match_score_pct), unsafe_allow_html=True)
        st.caption(f"📈 {match_delta}")
    
    with col2:
        st.markdown("""
        <div class="dashboard-metric-card">
            <div class="dashboard-metric-label">Est. Salary</div>
            <div class="dashboard-metric-value">HKD {0}k</div>
        </div>
        """.format(avg_salary // 1000), unsafe_allow_html=True)
        st.caption(f"💰 {salary_delta}")
    
    with col3:
        st.markdown("""
        <div class="dashboard-metric-card">
            <div class="dashboard-metric-label">Skill Gaps</div>
            <div class="dashboard-metric-value">{0}</div>
        </div>
        """.format(num_skill_gaps), unsafe_allow_html=True)
        st.caption(f"🎯 {gap_delta}")

def display_refine_results_section(matched_jobs, user_profile):
    """Display Refine Results section with filters"""
    st.markdown("---")
    with st.expander("🔧 Refine Results", expanded=False):
        st.markdown("### Adjust Search Criteria")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Target Domains
            current_domains = st.session_state.get('target_domains', [])
            target_domains = st.multiselect(
                "Target Domains (HK Focus)",
                options=["FinTech", "ESG & Sustainability", "Data Analytics", "Digital Transformation", 
                        "Investment Banking", "Consulting", "Technology", "Healthcare", "Education"],
                default=current_domains,
                key="refine_domains"
            )
        
        with col2:
            # Salary Expectations
            current_salary = st.session_state.get('salary_expectation', 0)
            salary_expectation = st.slider(
                "Min. Monthly Salary (HKD)",
                min_value=0,
                max_value=150000,
                value=current_salary,
                step=5000,
                help="Set to 0 to disable salary filtering",
                key="refine_salary"
            )
        
        force_refresh = st.checkbox(
            "Force new API fetch",
            value=False,
            help="Bypass cached results (only if results seem stale).",
            key="force_refresh_jobs_toggle"
        )
        
        if st.button("🔄 Apply Filters & Refresh", type="primary", use_container_width=True):
            # Update session state
            st.session_state.target_domains = target_domains
            st.session_state.salary_expectation = salary_expectation
            
            # Re-fetch and filter jobs
            search_query = " ".join(target_domains) if target_domains else "Hong Kong jobs"
            scraper = get_job_scraper()
            
            if scraper is None:
                st.error("⚠️ Job scraper not configured. Please check your RAPIDAPI_KEY in Streamlit secrets.")
                return
            
            with st.spinner("🔄 Refreshing results from Indeed..."):
                # Fetch jobs with cache
                # force_refresh is only True if user explicitly checked the checkbox
                # This is the ONLY place where force_refresh can be True (user-initiated)
                jobs = fetch_jobs_with_cache(
                    scraper,
                    search_query,
                    location="Hong Kong",
                    max_rows=25,
                    job_type="fulltime",
                    country="hk",
                    force_refresh=force_refresh  # Only True if user explicitly checked the box
                )
                
                if not jobs:
                    st.error("❌ No jobs found from Indeed. Please check your API configuration or try different search criteria.")
                    return
                
                # Show how many jobs were fetched before filtering
                total_fetched = len(jobs)
                
                # Apply domain filters
                if target_domains:
                    jobs = filter_jobs_by_domains(jobs, target_domains)
                
                # Apply salary filter
                if salary_expectation > 0:
                    jobs = filter_jobs_by_salary(jobs, salary_expectation)
                
                if not jobs:
                    st.warning(f"⚠️ No jobs match your filters. Found {total_fetched} jobs but none passed your criteria. Try reducing salary or selecting different domains.")
                    return
                
                # Re-index and search
                embedding_gen = get_embedding_generator()
                desired_matches = min(15, len(jobs))
                jobs_to_index_limit = _determine_index_limit(len(jobs), desired_matches)
                top_match_count = min(desired_matches, jobs_to_index_limit)
                search_engine = SemanticJobSearch(embedding_gen)
                search_engine.index_jobs(jobs, max_jobs_to_index=jobs_to_index_limit)
                
                # Use pre-computed resume embedding if available
                resume_embedding = st.session_state.get('resume_embedding')
                if not resume_embedding and st.session_state.resume_text:
                    resume_embedding = generate_and_store_resume_embedding(
                        st.session_state.resume_text,
                        st.session_state.user_profile if st.session_state.user_profile else None
                    )
                
                # Fallback: build query string if no resume embedding available
                resume_query = None
                if not resume_embedding:
                    if st.session_state.resume_text:
                        resume_query = st.session_state.resume_text
                        if st.session_state.user_profile.get('summary'):
                            profile_data = f"{st.session_state.user_profile.get('summary', '')} {st.session_state.user_profile.get('experience', '')} {st.session_state.user_profile.get('skills', '')}"
                            resume_query = f"{resume_query} {profile_data}"
                    else:
                        resume_query = f"{st.session_state.user_profile.get('summary', '')} {st.session_state.user_profile.get('experience', '')} {st.session_state.user_profile.get('skills', '')} {st.session_state.user_profile.get('education', '')}"
                
                results = search_engine.search(query=resume_query, top_k=top_match_count, resume_embedding=resume_embedding)
                
                # Calculate skill matches
                user_skills = st.session_state.user_profile.get('skills', '')
                for result in results:
                    job_skills = result['job'].get('skills', [])
                    skill_score, missing_skills = search_engine.calculate_skill_match(user_skills, job_skills)
                    result['skill_match_score'] = skill_score
                    result['missing_skills'] = missing_skills
                    
                    # Calculate combined match score (weighted: 60% semantic, 40% skill)
                    semantic_score = result.get('similarity_score', 0.0)
                    combined_score = (semantic_score * 0.6) + (skill_score * 0.4)
                    result['combined_match_score'] = combined_score
                
                # Sort results by combined match score (highest to lowest)
                results.sort(key=lambda x: x.get('combined_match_score', 0.0), reverse=True)
                
                st.session_state.matched_jobs = results
                st.session_state.dashboard_ready = True
                
                # Clean up memory after analysis
                gc.collect()
                
                st.rerun()


def display_ranked_matches_table(matched_jobs, user_profile):
    """Display Smart Ranked Matches Table with interactive dataframe
    
    Features:
    - Fixed Rank column (position 1, doesn't move when scrolling)
    - Click any row to expand and see full job details
    - Match score breakdown with detailed analysis
    - Rank is based on combined match score (60% semantic + 40% skill)
    """
    if not matched_jobs:
        return
    
    st.markdown("---")
    st.markdown("### Top AI-Ranked Opportunities")
    st.caption("💡 **Tip:** Click any row to expand and see full job description, match analysis, and application copilot")
    
    # Ensure all results have skill match scores calculated
    user_skills = user_profile.get('skills', '')
    
    # Helper function to calculate skill match
    def calc_skill_match(user_skills_str, job_skills_list):
        if not user_skills_str or not job_skills_list:
            return 0.0, []
        user_skills_lower = [s.lower().strip() for s in str(user_skills_str).split(',') if s.strip()]
        job_skills_lower = [s.lower().strip() for s in job_skills_list if isinstance(s, str) and s.strip()]
        if not user_skills_lower or not job_skills_lower:
            return 0.0, []
        matched_skills = []
        for job_skill in job_skills_lower:
            for user_skill in user_skills_lower:
                if job_skill in user_skill or user_skill in job_skill:
                    matched_skills.append(job_skill)
                    break
        match_score = len(matched_skills) / len(job_skills_lower) if job_skills_lower else 0.0
        missing_skills = [s for s in job_skills_lower if s not in matched_skills]
        return min(match_score, 1.0), missing_skills[:5]
    
    for result in matched_jobs:
        if 'skill_match_score' not in result:
            job_skills = result['job'].get('skills', [])
            skill_score, missing_skills = calc_skill_match(user_skills, job_skills)
            result['skill_match_score'] = skill_score
            result['missing_skills'] = missing_skills
        
        # Calculate combined match score if not already present
        if 'combined_match_score' not in result:
            semantic_score = result.get('similarity_score', 0.0)
            skill_score = result.get('skill_match_score', 0.0)
            combined_score = (semantic_score * 0.6) + (skill_score * 0.4)
            result['combined_match_score'] = combined_score
    
    # Sort matched_jobs by combined match score (highest to lowest)
    matched_jobs.sort(key=lambda x: x.get('combined_match_score', 0.0), reverse=True)
    
    # Create DataFrame
    table_data = []
    for i, result in enumerate(matched_jobs):
        job = result['job']
        semantic_score = result.get('similarity_score', 0.0)
        skill_score = result.get('skill_match_score', 0.0)
        # Use combined match score (60% semantic, 40% skill)
        match_score = result.get('combined_match_score', (semantic_score * 0.6) + (skill_score * 0.4))
        
        # Get key matching skills (first 3-4 skills from job that user has)
        job_skills = job.get('skills', [])
        matching_skills = []
        user_skills_list = [s.lower().strip() for s in str(user_skills).split(',') if s.strip()]
        for js in job_skills[:6]:
            if isinstance(js, str):
                js_lower = js.lower().strip()
                if any(us in js_lower or js_lower in us for us in user_skills_list):
                    matching_skills.append(js)
                    if len(matching_skills) >= 4:
                        break
        
        missing_critical = result.get('missing_skills', [])
        missing_critical_skill = missing_critical[0] if missing_critical else "None"
        
        table_data.append({
            'Rank': i + 1,  # Fixed rank index (1-based)
            'Match Score': int(match_score * 100),
            'Job Title': job['title'],
            'Company': job['company'],
            'Location': job['location'],
            'Key Matching Skills': matching_skills[:4] if matching_skills else [],
            'Missing Critical Skill': missing_critical_skill,
            '_index': i  # Internal index for selection
        })
    
    df = pd.DataFrame(table_data)
    
    # Configure column display
    column_config = {
        'Rank': st.column_config.NumberColumn(
            'Rank',
            help='Job ranking position (fixed, does not change when sorting)',
            width='small',
            format='%d'
        ),
        'Match Score': st.column_config.ProgressColumn(
            'Match Score',
            help='Combined match score: 60% semantic similarity + 40% skill overlap (jobs ranked by this)',
            min_value=0,
            max_value=100,
            format='%d%%'
        ),
        'Job Title': st.column_config.TextColumn(
            'Job Title',
            width='medium',
            help='Click to select and view full details'
        ),
        'Company': st.column_config.TextColumn(
            'Company',
            width='medium'
        ),
        'Location': st.column_config.TextColumn(
            'Location',
            width='small'
        ),
        'Key Matching Skills': st.column_config.ListColumn(
            'Key Matching Skills',
            help='Top skills you have that match this role'
        ),
        'Missing Critical Skill': st.column_config.TextColumn(
            'Missing Critical Skill',
            help='Most important skill gap for this role',
            width='medium'
        ),
        '_index': st.column_config.NumberColumn(
            '_index',
            width='small',
            help=None
        )
    }
    
    # Set column order with Rank first (fixed position - won't move when table scrolls)
    # Note: Rank is based on combined match score and remains fixed in position 1
    column_order = ['Rank', 'Match Score', 'Job Title', 'Company', 'Location', 'Key Matching Skills', 'Missing Critical Skill']
    
    # Reorder dataframe to ensure Rank is first
    df_display = df[column_order].copy()
    
    # Display dataframe with selection
    selected_rows = st.dataframe(
        df_display,
        column_config=column_config,
        hide_index=True,
        use_container_width=True,
        on_select="rerun",
        selection_mode="single-row"
    )
    
    # Store selected job index
    if selected_rows.selection.rows:
        selected_idx = df.iloc[selected_rows.selection.rows[0]]['_index']
        st.session_state.selected_job_index = int(selected_idx)
    else:
        st.session_state.selected_job_index = None

def display_match_breakdown(matched_jobs, user_profile):
    """Display Match Breakdown & Application Copilot in expander"""
    if st.session_state.selected_job_index is None:
        return
    
    selected_result = matched_jobs[st.session_state.selected_job_index]
    job = selected_result['job']
    semantic_score = selected_result.get('similarity_score', 0.0)
    skill_score = selected_result.get('skill_match_score', 0.0)
    missing_skills = selected_result.get('missing_skills', [])
    
    # Calculate skill overlap
    user_skills = user_profile.get('skills', '')
    job_skills = job.get('skills', [])
    user_skills_list = [s.lower().strip() for s in str(user_skills).split(',') if s.strip()]
    job_skills_list = [s.lower().strip() for s in job_skills if isinstance(s, str) and s.strip()]
    
    matched_skills_count = 0
    for js in job_skills_list:
        if any(us in js or js in us for us in user_skills_list):
            matched_skills_count += 1
    
    total_required = len(job_skills_list) if job_skills_list else 1
    skill_overlap_pct = (matched_skills_count / total_required * 100) if total_required > 0 else 0
    
    # Generate AI recruiter note
    text_gen = get_text_generator()
    if text_gen is None:
        recruiter_note = "AI analysis unavailable. Please configure Azure OpenAI credentials."
    else:
        recruiter_note = text_gen.generate_recruiter_note(job, user_profile, semantic_score, skill_score)
    
    # Get rank position (1-based)
    rank_position = st.session_state.selected_job_index + 1 if st.session_state.selected_job_index is not None else 0
    
    # Expander title with rank
    expander_title = f"📋 Rank #{rank_position}: {job['title']} at {job['company']}"
    
    with st.expander(expander_title, expanded=True):
        # Full Job Description Section
        st.markdown("#### 📝 Full Job Description")
        description_text = job.get('description', 'No description available.')
        if len(description_text) > 10000:
            st.info(f"📄 Full description ({len(description_text):,} characters)")
            st.text_area(
                "Job Description",
                value=description_text,
                height=400,
                key=f"job_desc_{st.session_state.selected_job_index}",
                label_visibility="collapsed"
            )
        else:
            st.markdown(description_text)
        
        st.markdown("---")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.markdown("#### 🎯 Match Analysis & Why This is a Fit")
            
            # Score breakdown with more detail
            st.markdown(f"""
            **Match Score Breakdown:**
            - **🎯 Skill Match Score (Ranking Factor):** {skill_score:.0%} ({matched_skills_count}/{total_required} skills matched)
              - This is the primary ranking factor. Jobs are sorted by this score.
            - **📊 Semantic Similarity Score:** {semantic_score:.0%}
              - Measures how well your experience contextually aligns with role requirements.
            - **⚖️ Combined Match Score:** {(semantic_score * 0.6 + skill_score * 0.4):.0%}
              - Weighted combination: 60% semantic + 40% skill overlap
            """)
            
            # Show matched skills
            if matched_skills_count > 0:
                matched_skills_display = []
                for js in job_skills_list:
                    if any(us in js or js in us for us in user_skills_list):
                        matched_skills_display.append(js)
                        if len(matched_skills_display) >= 10:
                            break
                if matched_skills_display:
                    st.success(f"✅ **Matched Skills:** {', '.join(matched_skills_display[:10])}")
            
            # Show missing skills
            if missing_skills:
                st.warning(f"⚠️ **Missing Skills:** {', '.join(missing_skills[:5])}")
            
            # Recruiter Note (AI-generated)
            st.markdown("---")
            st.info(f"**🤖 AI Recruiter Analysis:**\n\n{recruiter_note}")
        
        with col2:
            st.markdown("#### Application Copilot")
            
            # Accreditation Gap Action
            if missing_skills:
                top_missing = missing_skills[0]
                # Check if it's a certification-related skill
                cert_keywords = ['certification', 'certified', 'accreditation', 'license', 'pmp', 'scrum', 'hkicpa', 'cpa', 'cfa', 'cpa', 'aws', 'azure', 'gcp']
                is_cert = any(kw in top_missing.lower() for kw in cert_keywords)
                
                if is_cert:
                    st.warning(f"⚠️ **Crucial Gap:** This job highly values {top_missing}. Consider starting this certification.")
                else:
                    st.warning(f"⚠️ **Skill Gap:** Consider developing expertise in {top_missing}.")
            
            # Primary Action Button
            if st.button("✨ Tailor Resume for this Job", use_container_width=True, type="primary", key="tailor_resume_button"):
                st.session_state.selected_job = job
                st.session_state.show_resume_generator = True
                st.rerun()
            
            st.caption("Generates a citation-locked, AI-optimized CV emphasizing your matching skills.")
            
            # Apply to job link
            job_url = job.get('url', '#')
            if job_url and job_url != '#':
                st.markdown("---")
                st.link_button("🚀 Apply to Job", job_url, use_container_width=True, type="secondary")

def generate_pdf_from_json(resume_data, filename="resume.pdf"):
    """Generate a professional PDF file from structured resume JSON"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        # Create PDF in memory
        pdf_io = BytesIO()
        doc = SimpleDocTemplate(pdf_io, pagesize=letter,
                               rightMargin=0.75*inch, leftMargin=0.75*inch,
                               topMargin=0.5*inch, bottomMargin=0.5*inch)
        
        # Container for the 'Flowable' objects
        elements = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='black',
            spaceAfter=6,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            textColor='black',
            spaceAfter=6,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            textColor='black',
            spaceAfter=6,
            leading=12
        )
        
        contact_style = ParagraphStyle(
            'CustomContact',
            parent=styles['Normal'],
            fontSize=9,
            textColor='black',
            spaceAfter=12,
            alignment=TA_CENTER
        )
        
        # Header Section
        header = resume_data.get('header', {})
        if header.get('name'):
            elements.append(Paragraph(header['name'], title_style))
            elements.append(Spacer(1, 0.1*inch))
        
        # Contact information
        contact_info = []
        if header.get('email'):
            contact_info.append(header['email'])
        if header.get('phone'):
            contact_info.append(header['phone'])
        if header.get('location'):
            contact_info.append(header['location'])
        if header.get('linkedin'):
            contact_info.append(header['linkedin'])
        if header.get('portfolio'):
            contact_info.append(header['portfolio'])
        
        if contact_info:
            elements.append(Paragraph(' | '.join(contact_info), contact_style))
            elements.append(Spacer(1, 0.1*inch))
        
        # Professional Title
        if header.get('title'):
            elements.append(Paragraph(header['title'], contact_style))
            elements.append(Spacer(1, 0.15*inch))
        
        # Summary
        if resume_data.get('summary'):
            elements.append(Paragraph('Professional Summary', heading_style))
            elements.append(Paragraph(resume_data['summary'], normal_style))
            elements.append(Spacer(1, 0.1*inch))
        
        # Skills
        skills = resume_data.get('skills_highlighted', [])
        if skills:
            elements.append(Paragraph('Key Skills', heading_style))
            skills_text = ' • '.join(skills)
            elements.append(Paragraph(skills_text, normal_style))
            elements.append(Spacer(1, 0.1*inch))
        
        # Experience
        experience = resume_data.get('experience', [])
        if experience:
            elements.append(Paragraph('Professional Experience', heading_style))
            for exp in experience:
                # Company and Title
                exp_header_parts = []
                if exp.get('title'):
                    exp_header_parts.append(f"<b>{exp['title']}</b>")
                if exp.get('company'):
                    exp_header_parts.append(f" at {exp['company']}")
                if exp.get('dates'):
                    exp_header_parts.append(f" | <i>{exp['dates']}</i>")
                
                if exp_header_parts:
                    elements.append(Paragraph(''.join(exp_header_parts), normal_style))
                
                # Bullet points
                bullets = exp.get('bullets', [])
                for bullet in bullets:
                    if bullet.strip():
                        elements.append(Paragraph(f"• {bullet}", normal_style))
                
                elements.append(Spacer(1, 0.1*inch))
        
        # Education
        if resume_data.get('education'):
            elements.append(Paragraph('Education', heading_style))
            elements.append(Paragraph(resume_data['education'], normal_style))
            elements.append(Spacer(1, 0.1*inch))
        
        # Certifications
        if resume_data.get('certifications'):
            elements.append(Paragraph('Certifications & Awards', heading_style))
            elements.append(Paragraph(resume_data['certifications'], normal_style))
        
        # Build PDF
        doc.build(elements)
        pdf_io.seek(0)
        return pdf_io
        
    except Exception as e:
        st.error(f"Error generating PDF: {e}")
        return None

def format_resume_as_text(resume_data):
    """Format structured resume JSON as plain text"""
    text = []
    
    # Header
    header = resume_data.get('header', {})
    if header.get('name'):
        text.append(header['name'].upper())
        text.append("")
    
    # Contact info
    contact = []
    if header.get('email'):
        contact.append(header['email'])
    if header.get('phone'):
        contact.append(header['phone'])
    if header.get('location'):
        contact.append(header['location'])
    if header.get('linkedin'):
        contact.append(header['linkedin'])
    if header.get('portfolio'):
        contact.append(header['portfolio'])
    
    if contact:
        text.append(' | '.join(contact))
        text.append("")
    
    # Title
    if header.get('title'):
        text.append(header['title'])
        text.append("")
    
    # Summary
    if resume_data.get('summary'):
        text.append("PROFESSIONAL SUMMARY")
        text.append("-" * 50)
        text.append(resume_data['summary'])
        text.append("")
    
    # Skills
    skills = resume_data.get('skills_highlighted', [])
    if skills:
        text.append("KEY SKILLS")
        text.append("-" * 50)
        text.append(' • '.join(skills))
        text.append("")
    
    # Experience
    experience = resume_data.get('experience', [])
    if experience:
        text.append("PROFESSIONAL EXPERIENCE")
        text.append("-" * 50)
        for exp in experience:
            exp_line = exp.get('title', '')
            if exp.get('company'):
                exp_line += f" at {exp['company']}"
            if exp.get('dates'):
                exp_line += f" | {exp['dates']}"
            text.append(exp_line)
            
            bullets = exp.get('bullets', [])
            for bullet in bullets:
                if bullet.strip():
                    text.append(f"  • {bullet}")
            text.append("")
    
    # Education
    if resume_data.get('education'):
        text.append("EDUCATION")
        text.append("-" * 50)
        text.append(resume_data['education'])
        text.append("")
    
    # Certifications
    if resume_data.get('certifications'):
        text.append("CERTIFICATIONS & AWARDS")
        text.append("-" * 50)
        text.append(resume_data['certifications'])
    
    return '\n'.join(text)

def validate_secrets():
    """Validate that required secrets are configured. Returns True if valid, False otherwise."""
    try:
        required_secrets = ["AZURE_OPENAI_API_KEY", "AZURE_OPENAI_ENDPOINT", "RAPIDAPI_KEY"]
        missing_secrets = []
        
        for secret in required_secrets:
            if not st.secrets.get(secret):
                missing_secrets.append(secret)
        
        if missing_secrets:
            st.error(f"""
            ⚠️ **Missing Required Configuration**
            
            The following secrets are not configured in your Streamlit app:
            - {', '.join(missing_secrets)}
            
            Please configure these in your Streamlit Cloud secrets or local `.streamlit/secrets.toml` file.
            See `.streamlit/secrets.toml.example` for the required format.
            """)
            return False
        
        return True
    except Exception as e:
        st.error(f"⚠️ Error validating secrets: {e}")
        return False

def render_hero_banner(user_profile, matched_jobs=None):
    """Render the Modern Hero banner with personalized welcome message and Base64 logo watermark"""
    # Get user name or use default
    user_name = user_profile.get('name', '') if user_profile else ''
    if not user_name or user_name == 'N/A':
        user_name = 'Professional'
    
    # Generate AI-insights focused subtitle based on analysis state
    if matched_jobs and len(matched_jobs) > 0:
        # Calculate average match score
        avg_score = sum(r.get('combined_match_score', 0) for r in matched_jobs) / len(matched_jobs)
        subtitle = f"Your AI-powered career analysis is ready. We found {len(matched_jobs)} matching opportunities with {int(avg_score * 100)}% average fit."
    elif user_profile and user_profile.get('skills'):
        subtitle = "Your profile is loaded. Unlock AI-powered insights to discover your market positioning and best-fit opportunities."
    else:
        subtitle = "Leverage AI-powered insights to discover your market value, skill gaps, and best-fit opportunities in Hong Kong."
    
    # Render Modern Hero banner with Base64 logo watermark
    st.markdown(f"""
    <div class="hero-container">
        {_logo_html}
        <div class="hero-content">
            <div style="color: var(--cyan); font-weight: 600; margin-bottom: 8px; font-size: 12px; letter-spacing: 1.5px;">CAREERLENS AI INTELLIGENCE</div>
            <div class="hero-title">Welcome back, {user_name}.</div>
            <div class="hero-subtitle">{subtitle}</div>
        </div>
    </div>
    """, unsafe_allow_html=True)


def main():
    try:
        # Check if resume generator should be shown
        if st.session_state.get('show_resume_generator', False):
            display_resume_generator()
            return
        
        # Render sidebar with controls
        render_sidebar()
        
        # Render hero banner at the top of main content
        render_hero_banner(
            st.session_state.user_profile if st.session_state.user_profile else {},
            st.session_state.matched_jobs if st.session_state.get('dashboard_ready', False) else None
        )
        
        # Main dashboard area - only show after analysis
        if not st.session_state.get('dashboard_ready', False) or not st.session_state.matched_jobs:
            # Show empty state (but hero banner is already shown above)
            st.info("👆 Upload your CV in the sidebar and click 'Analyze Profile & Find Matches' to see your market positioning and ranked opportunities.")
            return
        
        # Display Market Positioning Profile (Top Section)
        display_market_positioning_profile(
            st.session_state.matched_jobs,
            st.session_state.user_profile
        )
        
        # Display Refine Results Section
        display_refine_results_section(
            st.session_state.matched_jobs,
            st.session_state.user_profile
        )
        
        # Display Smart Ranked Matches Table (Middle Section)
        display_ranked_matches_table(
            st.session_state.matched_jobs,
            st.session_state.user_profile
        )
        
        # Display Match Breakdown & Application Copilot (Bottom Section)
        display_match_breakdown(
            st.session_state.matched_jobs,
            st.session_state.user_profile
        )
    except Exception as e:
        st.error(f"""
        ❌ **Application Error**
        
        An unexpected error occurred: {e}
        
        Please check:
        1. All required secrets are configured
        2. All dependencies are installed
        3. The application logs for more details
        """)
        st.exception(e)

def _check_startup_health():
    """Perform startup health checks to ensure the app can run.
    
    Returns True if all checks pass, False otherwise.
    """
    issues = []
    
    # Check required imports
    if not PYPDF2_AVAILABLE:
        issues.append("PyPDF2 not available - PDF upload will be disabled")
    if not DOCX_AVAILABLE:
        issues.append("python-docx not available - DOCX upload will be disabled")
    if not CHROMADB_AVAILABLE:
        issues.append("ChromaDB not available - using simple in-memory storage")
    
    # Check secrets (gracefully)
    try:
        secrets_available = hasattr(st, 'secrets')
        if secrets_available:
            azure_key = st.secrets.get("AZURE_OPENAI_API_KEY")
            azure_endpoint = st.secrets.get("AZURE_OPENAI_ENDPOINT")
            rapidapi_key = st.secrets.get("RAPIDAPI_KEY")
            
            if not azure_key or not azure_endpoint:
                issues.append("Azure OpenAI credentials not configured")
            if not rapidapi_key:
                issues.append("RapidAPI key not configured")
    except Exception:
        # Secrets not available yet (first load) - this is okay
        pass
    
    # Log issues but don't block startup
    if issues:
        for issue in issues:
            st.sidebar.warning(f"⚠️ {issue}")
    
    return True  # Always allow startup, show warnings instead


if __name__ == "__main__":
    # Run garbage collection before startup
    gc.collect()
    
    # Perform startup health checks (non-blocking)
    _check_startup_health()
    
    # Wrap main() in error handling to prevent crashes
    try:
        main()
    except MemoryError:
        # Handle memory errors gracefully
        gc.collect()
        st.error("""
        ❌ **Memory Error**
        
        The application ran out of memory. This can happen on Streamlit Cloud with large datasets.
        
        **Solutions:**
        1. Refresh the page to clear cached data
        2. Try analyzing fewer jobs at once
        3. Close other browser tabs to free memory
        """)
    except Exception as e:
        st.error(f"""
        ❌ **Startup Error**
        
        The application failed to start: {e}
        
        This is likely due to:
        1. Missing or incorrect secrets configuration
        2. Missing dependencies
        3. A code error in the application
        
        Please check your Streamlit Cloud logs for more details.
        """)
        st.exception(e)
    finally:
        # Final garbage collection
        gc.collect()
